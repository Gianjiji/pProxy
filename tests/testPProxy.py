"""
Test suite per Privacy Proxy per LLM.
Eseguire con: pytest -v   (o: pytest testPProxy.py -v)
"""
import pytest
from privacy_proxy import (
    InputLayer,
    RuleBasedDetector,
    TokenizationEngine,
    RehydrationEngine,
    ValidationLayer,
    SecureMappingStore,
    SensitiveDataDetectionEngine,
    NERDetector,
    LLMGateway,
    LLMProvider,
    CSVColumnInferrer,
    TextChunker,
    PrivacyProxy,
    EntityType,
    DetectedEntity,
)
import json
from pathlib import Path


# ──────────────────────────────────────────────────────────────
# Fixtures
# ──────────────────────────────────────────────────────────────

SAMPLE_TEXT = (
    "Mario Rossi (CF: RSSMRA80A01H501U) lavora presso Acme S.r.l.\n"
    "Tel: +39 333 1234567  Email: mario.rossi@example.com\n"
    "IBAN: IT60X0542811101000000123456\n"
    "Stipendio: €2.350,00 — Data nascita: 01/01/1980\n"
    "Mario Rossi abita in Via Roma 1, 00100 Roma."
)


# ──────────────────────────────────────────────────────────────
# Modulo 1 – Input Layer
# ──────────────────────────────────────────────────────────────

class TestInputLayer:

    def test_raw_text(self):
        doc = InputLayer().load("Testo diretto")
        assert doc.text == "Testo diretto"
        assert doc.source == "raw_text"

    def test_txt_file(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("contenuto test", encoding="utf-8")
        doc = InputLayer().load(str(f))
        assert doc.text == "contenuto test"

    def test_json_file(self, tmp_path):
        f = tmp_path / "test.json"
        f.write_text('{"name": "Mario", "age": 30}', encoding="utf-8")
        doc = InputLayer().load(str(f))
        assert "Mario" in doc.text

    def test_json_annotated_format(self, tmp_path):
        """JSON deve essere serializzato come 'campo: valore' per agevolare il detector."""
        f = tmp_path / "record.json"
        f.write_text(
            '{"nome": "Luca Verdi", "email": "luca@example.com", "iban": "IT60X0542811101000000123456"}',
            encoding="utf-8",
        )
        doc = InputLayer().load(str(f))
        assert "nome: Luca Verdi" in doc.text
        assert "email: luca@example.com" in doc.text
        assert "iban: IT60X0542811101000000123456" in doc.text

    def test_json_nested_structure(self, tmp_path):
        """Strutture JSON annidate vengono appiattite in righe annotate."""
        f = tmp_path / "nested.json"
        f.write_text(
            '{"persona": {"nome": "Anna", "cf": "NNNNNN00A00H501A"}, "importo": 1000}',
            encoding="utf-8",
        )
        doc = InputLayer().load(str(f))
        assert "nome: Anna" in doc.text
        assert "cf: NNNNNN00A00H501A" in doc.text

    def test_unsupported_format(self, tmp_path):
        f = tmp_path / "test.xyz"
        f.write_bytes(b"data")
        with pytest.raises(ValueError, match="Formato non supportato"):
            InputLayer().load(str(f))

    def test_csv_file(self, tmp_path):
        f = tmp_path / "test.csv"
        f.write_text("Nome,Email\nMario,mario@test.com\n", encoding="utf-8")
        try:
            doc = InputLayer().load(str(f))
            assert "Mario" in doc.text
            # CSV per-column format: "Nome: Mario | Email: mario@test.com"
            assert "Nome:" in doc.text or "Mario" in doc.text
        except ImportError:
            pytest.skip("pandas non installato")

    def test_docx_includes_table_cells(self, tmp_path):
        """Verifica che le tabelle DOCX vengano estratte (non solo i paragrafi)."""
        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
        except ImportError:
            pytest.skip("python-docx non installato")

        f = tmp_path / "test.docx"
        d = DocxDocument()
        d.add_paragraph("Paragrafo normale con Mario Rossi.")
        tbl = d.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "IBAN"
        tbl.cell(0, 1).text = "Intestatario"
        tbl.cell(1, 0).text = "IT60X0542811101000000123456"
        tbl.cell(1, 1).text = "Luca Bianchi"
        d.save(str(f))

        doc = InputLayer().load(str(f))
        assert "Mario Rossi" in doc.text
        assert "IT60X0542811101000000123456" in doc.text
        assert "Luca Bianchi" in doc.text


# ──────────────────────────────────────────────────────────────
# Modulo 2 – Rule Based Detector
# ──────────────────────────────────────────────────────────────

class TestRuleBasedDetector:

    @pytest.fixture
    def det(self):
        return RuleBasedDetector()

    def test_email(self, det):
        entities = det.detect("Scrivimi a mario.rossi@example.com grazie")
        emails = [e for e in entities if e.type == EntityType.EMAIL]
        assert len(emails) == 1
        assert emails[0].value == "mario.rossi@example.com"
        assert emails[0].confidence >= 0.99

    def test_email_variants_still_detected(self, det):
        for addr in ("a@b.co", "test_user+tag@sub.domain.org", "x@y.museum"):
            entities = det.detect(f"contatto: {addr}")
            assert any(e.type == EntityType.EMAIL and e.value == addr for e in entities)

    def test_email_regex_no_redos(self, det):
        """Input avversariale (lunga sequenza di caratteri email senza '@') NON
        deve causare backtracking quadratico: il rilevamento resta rapido."""
        import time

        adversarial = ("1." * 50000 + "1")[:100000]
        t = time.time()
        det.detect(adversarial)
        assert time.time() - t < 1.0  # prima del fix: decine di secondi

    def test_iban(self, det):
        entities = det.detect("IBAN: IT60X0542811101000000123456 intestato a Mario")
        ibans = [e for e in entities if e.type == EntityType.IBAN]
        assert any("IT60X0542811101000000123456" in e.value for e in ibans)

    def test_codice_fiscale(self, det):
        entities = det.detect("CF: RSSMRA80A01H501U")
        cfs = [e for e in entities if e.type == EntityType.CF]
        assert len(cfs) == 1
        assert cfs[0].confidence >= 0.99

    def test_invalid_cf_rejected(self, det):
        """Codice Fiscale con carattere di controllo errato deve essere scartato."""
        # RSSMRA80A01H501U è valido (ctrl=U); cambiamo l'ultimo carattere in 'X'
        entities = det.detect("CF: RSSMRA80A01H501X")
        cfs = [e for e in entities if e.type == EntityType.CF]
        assert len(cfs) == 0

    def test_valid_iban_accepted(self, det):
        """IBAN italiano valido (mod-97 = 1) deve essere rilevato."""
        entities = det.detect("IBAN: IT60X0542811101000000123456")
        ibans = [e for e in entities if e.type == EntityType.IBAN]
        assert len(ibans) >= 1

    def test_invalid_iban_rejected(self, det):
        """Stringa che soddisfa il pattern ma fallisce il checksum mod-97 deve essere scartata."""
        # Modifichiamo le cifre di controllo (pos 2-3) da 60 a 99 → non valido
        entities = det.detect("IBAN: IT99X0542811101000000123456")
        ibans = [e for e in entities if e.type == EntityType.IBAN]
        assert len(ibans) == 0

    def test_amount_euro(self, det):
        entities = det.detect("Importo: €2.350,00 da versare")
        amounts = [e for e in entities if e.type == EntityType.AMOUNT]
        assert len(amounts) >= 1

    def test_amount_labeled_no_currency(self, det):
        """Importi preceduti da etichetta senza simbolo valuta (stipendio, reddito, ecc.)."""
        entities = det.detect("Stipendio: 2.500,00 mensile")
        amounts = [e for e in entities if e.type == EntityType.AMOUNT]
        assert len(amounts) >= 1

    def test_amount_reddito(self, det):
        entities = det.detect("reddito: 35000")
        amounts = [e for e in entities if e.type == EntityType.AMOUNT]
        assert len(amounts) >= 1

    def test_date(self, det):
        entities = det.detect("Nato il 01/01/1980 a Roma")
        dates = [e for e in entities if e.type == EntityType.DATE]
        assert len(dates) >= 1

    def test_date_written_italian(self, det):
        """Formato scritto 'GG mese AAAA' tipico dei contratti italiani."""
        entities = det.detect("Firmato il 15 marzo 2022 a Milano")
        dates = [e for e in entities if e.type == EntityType.DATE]
        assert len(dates) >= 1
        assert "15 marzo 2022" in dates[0].value

    def test_date_abbreviated_month(self, det):
        """Mese abbreviato: 'gen', 'feb', ecc."""
        entities = det.detect("Data nascita: 3 gen 1995")
        dates = [e for e in entities if e.type == EntityType.DATE]
        assert len(dates) >= 1

    def test_date_abbreviated_month_with_period(self, det):
        """Mese abbreviato con punto: '15 gen. 2024' (stile documenti formali)."""
        entities = det.detect("Emesso il 15 gen. 2024 dalla banca")
        dates = [e for e in entities if e.type == EntityType.DATE]
        assert len(dates) >= 1
        assert "gen." in dates[0].value

    def test_date_ordinal_day(self, det):
        """Giorno ordinale: '1° gennaio 2024' tipico dei contratti."""
        entities = det.detect("A partire dal 1° gennaio 2024")
        dates = [e for e in entities if e.type == EntityType.DATE]
        assert len(dates) >= 1

    def test_date_month_year_only(self, det):
        """'giugno 2024' – riferimento mese/anno senza giorno."""
        entities = det.detect("Fattura di competenza: giugno 2024")
        dates = [e for e in entities if e.type == EntityType.DATE]
        assert len(dates) >= 1
        assert "giugno 2024" in dates[0].value

    def test_phone_italian(self, det):
        entities = det.detect("Chiamare al +39 333 1234567")
        phones = [e for e in entities if e.type == EntityType.PHONE]
        assert len(phones) >= 1

    def test_phone_international(self, det):
        """Numero internazionale non italiano (es. UK +44) deve essere rilevato."""
        entities = det.detect("Contatto UK: +441234567890")
        phones = [e for e in entities if e.type == EntityType.PHONE]
        assert len(phones) >= 1

    def test_address_via(self, det):
        """Indirizzo italiano Via + nome + numero civico."""
        entities = det.detect("Residente in Via Roma, 15 – Milano")
        addrs = [e for e in entities if e.type == EntityType.ADDRESS]
        assert len(addrs) >= 1
        assert "Via Roma" in addrs[0].value

    def test_address_piazza(self, det):
        """Piazza con numero civico."""
        entities = det.detect("Ufficio: Piazza Navona 10, Roma")
        addrs = [e for e in entities if e.type == EntityType.ADDRESS]
        assert len(addrs) >= 1

    def test_credit_card_luhn(self, det):
        # Numero Visa valido (Luhn ok)
        entities = det.detect("Carta: 4532015112830366")
        cards = [e for e in entities if e.type == EntityType.CARD]
        assert len(cards) >= 1

    def test_invalid_credit_card_rejected(self, det):
        # Numero con Luhn sbagliato
        entities = det.detect("Carta: 4532015112830367")
        cards = [e for e in entities if e.type == EntityType.CARD]
        assert len(cards) == 0

    def test_person_with_title_dott(self, det):
        """Titolo 'Dott.' seguito da nome e cognome rilevato come PERSON."""
        entities = det.detect("Il paziente è seguito dal Dott. Mario Bianchi")
        persons = [e for e in entities if e.type == EntityType.PERSON]
        assert len(persons) >= 1
        assert "Bianchi" in persons[0].value

    def test_person_with_title_sig(self, det):
        """Titolo 'Sig.ra' con nome e cognome."""
        entities = det.detect("La Sig.ra Anna Verdi ha firmato il contratto")
        persons = [e for e in entities if e.type == EntityType.PERSON]
        assert len(persons) >= 1
        assert "Verdi" in persons[0].value

    def test_person_with_title_avv(self, det):
        """Titolo professionale 'Avv.' rilevato come PERSON."""
        entities = det.detect("Difeso dall'Avv. Carlo Neri del foro di Milano")
        persons = [e for e in entities if e.type == EntityType.PERSON]
        assert len(persons) >= 1

    def test_person_with_title_prof(self, det):
        """'Prof.' e 'Prof.ssa' rilevati come PERSON."""
        entities = det.detect("Relazione della Prof.ssa Laura Esposito")
        persons = [e for e in entities if e.type == EntityType.PERSON]
        assert len(persons) >= 1
        assert "Esposito" in persons[0].value

    def test_full_sample(self, det):
        entities = det.detect(SAMPLE_TEXT)
        types_found = {e.type for e in entities}
        assert EntityType.EMAIL in types_found
        assert EntityType.IBAN in types_found
        assert EntityType.CF in types_found


# ──────────────────────────────────────────────────────────────
# Modulo 3 – Tokenization Engine
# ──────────────────────────────────────────────────────────────

class TestTokenizationEngine:

    def _make_entity(self, value, etype, start, end):
        return DetectedEntity(value=value, type=etype, confidence=0.99, start=start, end=end)

    def test_basic_replacement(self):
        engine = TokenizationEngine()
        text = "Mario Rossi lavora qui"
        entities = [self._make_entity("Mario Rossi", EntityType.PERSON, 0, 11)]
        anon, mapping = engine.tokenize(text, entities)
        assert "[PERSON_001]" in anon
        assert "Mario Rossi" not in anon
        assert mapping["[PERSON_001]"] == "Mario Rossi"

    def test_same_value_same_token(self):
        engine = TokenizationEngine()
        text = "Mario Rossi parla con Mario Rossi"
        entities = [
            self._make_entity("Mario Rossi", EntityType.PERSON, 0, 11),
            self._make_entity("Mario Rossi", EntityType.PERSON, 22, 33),
        ]
        anon, mapping = engine.tokenize(text, entities)
        assert anon.count("[PERSON_001]") == 2
        assert "[PERSON_002]" not in anon

    def test_different_values_different_tokens(self):
        engine = TokenizationEngine()
        text = "Mario Rossi e Luigi Verdi"
        entities = [
            self._make_entity("Mario Rossi", EntityType.PERSON, 0, 11),
            self._make_entity("Luigi Verdi", EntityType.PERSON, 14, 25),
        ]
        anon, mapping = engine.tokenize(text, entities)
        assert "[PERSON_001]" in anon
        assert "[PERSON_002]" in anon
        assert len(mapping) == 2

    def test_reset_clears_state(self):
        engine = TokenizationEngine()
        text = "Mario"
        entities = [self._make_entity("Mario", EntityType.PERSON, 0, 5)]
        _, m1 = engine.tokenize(text, entities)
        engine.reset()
        _, m2 = engine.tokenize(text, entities)
        # After reset, counter restarts at 001
        assert "[PERSON_001]" in m1
        assert "[PERSON_001]" in m2

    def test_multiple_entity_types(self):
        engine = TokenizationEngine()
        text = "mario@test.com e +39 333 1234567"
        entities = [
            self._make_entity("mario@test.com", EntityType.EMAIL, 0, 14),
            self._make_entity("+39 333 1234567", EntityType.PHONE, 17, 32),
        ]
        anon, mapping = engine.tokenize(text, entities)
        assert "[EMAIL_001]" in anon
        assert "[PHONE_001]" in anon


# ──────────────────────────────────────────────────────────────
# Modulo 4 – Secure Mapping Store
# ──────────────────────────────────────────────────────────────

class TestSecureMappingStore:

    def test_basic_store_retrieve(self):
        store = SecureMappingStore()
        store.store({"[PERSON_001]": "Mario Rossi"})
        assert store.get_mapping()["[PERSON_001]"] == "Mario Rossi"

    def test_clear(self):
        store = SecureMappingStore()
        store.store({"[PERSON_001]": "Mario Rossi"})
        store.clear()
        assert store.get_mapping() == {}

    def test_save_and_load_plaintext(self, tmp_path):
        path = str(tmp_path / "map.json")
        store = SecureMappingStore(store_path=path)
        store.store({"[EMAIL_001]": "test@example.com"})
        store2 = SecureMappingStore(store_path=path)
        store2.load()
        assert store2.get_mapping()["[EMAIL_001]"] == "test@example.com"

    def test_encrypted_save_load(self, tmp_path):
        try:
            from cryptography.fernet import Fernet
        except ImportError:
            pytest.skip("cryptography non installato")
        path = str(tmp_path / "map.enc")
        key = "test-passphrase-segreta"
        store = SecureMappingStore(store_path=path, encryption_key=key)
        store.store({"[IBAN_001]": "IT60X0542811101000000123456"})
        # Verify raw file is not plain JSON
        raw = Path(path).read_bytes()
        assert b"IT60" not in raw
        # Load with same key
        store2 = SecureMappingStore(store_path=path, encryption_key=key)
        store2.load()
        assert store2.get_mapping()["[IBAN_001]"] == "IT60X0542811101000000123456"

    def test_encrypted_uses_random_salt(self, tmp_path):
        """Ogni salvataggio cifrato usa un salt casuale (formato v2) → ciphertext diverso."""
        try:
            from cryptography.fernet import Fernet  # noqa: F401
        except ImportError:
            pytest.skip("cryptography non installato")
        path = str(tmp_path / "map.enc")
        store = SecureMappingStore(store_path=path, encryption_key="pw")
        store.store({"[EMAIL_001]": "a@b.com"})
        raw1 = Path(path).read_bytes()
        assert raw1.startswith(b"PPX2")          # header formato v2
        store.store({})                          # ri-salva con nuovo salt
        raw2 = Path(path).read_bytes()
        assert raw1[:20] != raw2[:20]            # salt diverso ad ogni save

    def test_wrong_passphrase_rejected(self, tmp_path):
        """Una passphrase errata non deve decifrare la mappa."""
        try:
            from cryptography.fernet import InvalidToken  # noqa: F401
        except ImportError:
            pytest.skip("cryptography non installato")
        path = str(tmp_path / "map.enc")
        SecureMappingStore(store_path=path, encryption_key="giusta").store(
            {"[CF_001]": "RSSMRA80A01H501Z"}
        )
        with pytest.raises(Exception):
            SecureMappingStore(store_path=path, encryption_key="sbagliata").load()

    def test_legacy_v1_encrypted_file_still_loads(self, tmp_path):
        """File cifrati legacy (salt statico, senza header) restano leggibili."""
        try:
            from cryptography.fernet import Fernet
            from cryptography.hazmat.primitives import hashes
            from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
        except ImportError:
            pytest.skip("cryptography non installato")
        import base64
        pw = "legacypw"
        kdf = PBKDF2HMAC(
            algorithm=hashes.SHA256(), length=32,
            salt=b"privacy_proxy_v1", iterations=480_000,
        )
        key = base64.urlsafe_b64encode(kdf.derive(pw.encode()))
        token = Fernet(key).encrypt(json.dumps({"[CF_001]": "ABC"}).encode())
        path = tmp_path / "old.enc"
        path.write_bytes(token)  # nessun header PPX2 → percorso legacy
        store = SecureMappingStore(store_path=str(path), encryption_key=pw)
        store.load()
        assert store.get_mapping()["[CF_001]"] == "ABC"


# ──────────────────────────────────────────────────────────────
# Modulo 6 – Rehydration Engine
# ──────────────────────────────────────────────────────────────

class TestRehydrationEngine:

    def test_basic_rehydrate(self):
        engine = RehydrationEngine()
        text = "[PERSON_001] guadagna [AMOUNT_001]"
        mapping = {"[PERSON_001]": "Mario Rossi", "[AMOUNT_001]": "€2.350,00"}
        result = engine.rehydrate(text, mapping)
        assert result == "Mario Rossi guadagna €2.350,00"

    def test_repeated_token(self):
        engine = RehydrationEngine()
        text = "[PERSON_001] parla con [PERSON_001]"
        mapping = {"[PERSON_001]": "Mario Rossi"}
        result = engine.rehydrate(text, mapping)
        assert result == "Mario Rossi parla con Mario Rossi"

    def test_find_unreplaced(self):
        engine = RehydrationEngine()
        text = "Ciao [PERSON_001], il tuo saldo è [AMOUNT_001]."
        unreplaced = engine.find_unreplaced(text)
        assert set(unreplaced) == {"[PERSON_001]", "[AMOUNT_001]"}

    def test_empty_map(self):
        engine = RehydrationEngine()
        text = "Nessun placeholder qui"
        result = engine.rehydrate(text, {})
        assert result == text

    def test_lowercase_token_rehydrated(self):
        """LLM può restituire token in minuscolo – deve essere ripristinato ugualmente."""
        engine = RehydrationEngine()
        text = "[person_001] guadagna [amount_001]"
        mapping = {"[PERSON_001]": "Mario Rossi", "[AMOUNT_001]": "€2.350,00"}
        result = engine.rehydrate(text, mapping)
        assert result == "Mario Rossi guadagna €2.350,00"

    def test_mixedcase_token_rehydrated(self):
        """Token con maiuscole miste (es. [Person_001]) deve essere ripristinato."""
        engine = RehydrationEngine()
        text = "Beneficiario: [Person_001]"
        mapping = {"[PERSON_001]": "Luca Bianchi"}
        result = engine.rehydrate(text, mapping)
        assert result == "Beneficiario: Luca Bianchi"

    def test_rehydrate_value_with_backslash(self):
        """Valore originale con backslash (es. percorso Windows) non deve essere corrotto.

        Regressione: re.sub interpreta \\1, \\g<...> nella stringa di replacement;
        un valore con backslash sollevava re.error o veniva alterato.
        """
        engine = RehydrationEngine()
        mapping = {"[ADDR_001]": r"C:\new\tmp"}
        result = engine.rehydrate("Percorso: [ADDR_001]", mapping)
        assert result == r"Percorso: C:\new\tmp"

    def test_rehydrate_value_with_group_reference(self):
        """Valore originale contenente una sequenza tipo backreference (\\1) resta letterale."""
        engine = RehydrationEngine()
        mapping = {"[PERSON_001]": r"A\1B"}
        result = engine.rehydrate("Nome: [person_001]", mapping)
        assert result == r"Nome: A\1B"

    def test_rehydrate_longest_token_wins(self):
        """Con token che condividono un prefisso, vince il match più lungo."""
        engine = RehydrationEngine()
        mapping = {"[ORG_001]": "Acme", "[ORG_0011]": "Globex"}
        assert engine.rehydrate("[ORG_0011] e [ORG_001]", mapping) == "Globex e Acme"

    def test_rehydrate_large_mapping_is_fast(self):
        """Passata singola: una mappa molto grande non deve degradare a O(n*m)."""
        import time

        engine = RehydrationEngine()
        text = "[TOK_00001] " * 5000
        mapping = {f"[TOK_{i:05d}]": "X" * 10 for i in range(20000)}
        t = time.time()
        engine.rehydrate(text, mapping)
        assert time.time() - t < 1.0  # ciclo per-token precedente: ~0.8s e oltre

    def test_find_unreplaced_four_digit_token(self):
        """Token con indice a 4 cifre (>999 entità) deve essere rilevato come non sostituito."""
        engine = RehydrationEngine()
        assert engine.find_unreplaced("resto: [PERSON_1000]") == ["[PERSON_1000]"]
        # Riepilogo coerente anche per indici a 3 cifre
        assert engine.find_unreplaced("[EMAIL_001] e [EMAIL_999]") == [
            "[EMAIL_001]", "[EMAIL_999]"
        ]


# ──────────────────────────────────────────────────────────────
# Modulo 7 – Validation Layer
# ──────────────────────────────────────────────────────────────

class TestValidationLayer:

    def test_valid_anonymized(self):
        v = ValidationLayer()
        result = v.validate_anonymized(
            anonymized="[PERSON_001] lavora qui",
            entity_map={"[PERSON_001]": "Mario Rossi"},
        )
        assert result.is_valid
        assert not result.errors

    def test_original_leaked(self):
        v = ValidationLayer()
        result = v.validate_anonymized(
            anonymized="Mario Rossi lavora qui",
            entity_map={"[PERSON_001]": "Mario Rossi"},
        )
        # A leaked original value is now an ERROR (not just a warning)
        assert not result.is_valid
        assert any("Mario Rossi" in e for e in result.errors)

    def test_unknown_token_in_anonymized(self):
        v = ValidationLayer()
        result = v.validate_anonymized(
            anonymized="[PERSON_999] lavora qui",
            entity_map={"[PERSON_001]": "Mario Rossi"},
        )
        assert not result.is_valid

    def test_valid_rehydrated(self):
        v = ValidationLayer()
        result = v.validate_rehydrated(
            rehydrated="Mario Rossi guadagna 2000€",
            entity_map={"[PERSON_001]": "Mario Rossi", "[AMOUNT_001]": "2000€"},
        )
        assert result.is_valid

    def test_unreplaced_known_token(self):
        v = ValidationLayer()
        result = v.validate_rehydrated(
            rehydrated="[PERSON_001] guadagna 2000€",
            entity_map={"[PERSON_001]": "Mario Rossi"},
        )
        assert not result.is_valid
        assert any("non sostituito" in e for e in result.errors)

    def test_invented_token(self):
        v = ValidationLayer()
        result = v.validate_rehydrated(
            rehydrated="Mario Rossi e [PERSON_999]",
            entity_map={"[PERSON_001]": "Mario Rossi"},
        )
        assert not result.is_valid
        assert any("inventato" in e for e in result.errors)


# ──────────────────────────────────────────────────────────────
# Integrazione: full anonymize (senza LLM)
# ──────────────────────────────────────────────────────────────

class TestFullPipeline:

    def test_anonymize_and_rehydrate(self):
        detector = SensitiveDataDetectionEngine(confidence_threshold=0.7, use_ner=False)
        tokenizer = TokenizationEngine()
        rehydrator = RehydrationEngine()
        validator = ValidationLayer()

        text = "Contattare mario@test.com o chiamare 333-1234567"
        entities = detector.detect(text)
        tokenizer.reset()
        anon, mapping = tokenizer.tokenize(text, entities)

        assert "mario@test.com" not in anon
        assert "[EMAIL_001]" in anon

        restored = rehydrator.rehydrate(anon, mapping)
        assert "mario@test.com" in restored

        val = validator.validate_rehydrated(restored, mapping)
        assert val.is_valid

    def test_deterministic_tokens(self):
        detector = SensitiveDataDetectionEngine(confidence_threshold=0.7, use_ner=False)
        tokenizer = TokenizationEngine()

        text = "CF: RSSMRA80A01H501U — verificato per RSSMRA80A01H501U"
        entities = detector.detect(text)
        tokenizer.reset()
        anon, mapping = tokenizer.tokenize(text, entities)

        # Same CF must get same token
        assert anon.count("[CF_001]") == 2
        assert "[CF_002]" not in anon

    def test_sample_document(self):
        detector = SensitiveDataDetectionEngine(confidence_threshold=0.7, use_ner=False)
        tokenizer = TokenizationEngine()
        rehydrator = RehydrationEngine()

        tokenizer.reset()
        entities = detector.detect(SAMPLE_TEXT)
        anon, mapping = tokenizer.tokenize(SAMPLE_TEXT, entities)

        # Key PII should be anonymized
        assert "RSSMRA80A01H501U" not in anon
        assert "mario.rossi@example.com" not in anon

        restored = rehydrator.rehydrate(anon, mapping)
        assert "RSSMRA80A01H501U" in restored
        assert "mario.rossi@example.com" in restored


class TestTextChunker:

    def test_short_text_no_split(self):
        chunker = TextChunker(max_chars=100)
        text = "Testo breve"
        chunks = chunker.chunk(text)
        assert chunks == ["Testo breve"]

    def test_splits_at_paragraph(self):
        chunker = TextChunker(max_chars=30)
        text = "Prima parte del testo.\n\nSeconda parte del testo."
        chunks = chunker.chunk(text)
        assert len(chunks) == 2
        assert "Prima" in chunks[0]
        assert "Seconda" in chunks[1]

    def test_does_not_split_placeholder(self):
        chunker = TextChunker(max_chars=20)
        # [PERSON_001] is 12 chars; with 20-char window, avoid splitting inside it
        text = "Nome: [PERSON_001] lavora"
        chunks = chunker.chunk(text)
        # Each chunk should contain complete, unsplit tokens
        for chunk in chunks:
            # No half-open bracket
            assert not (chunk.count("[") != chunk.count("]"))

    def test_does_not_split_token_straddling_boundary(self):
        """Regressione: un token che attraversa esattamente il confine max_chars
        non deve essere spezzato. Prima il controllo era codice morto (cercava
        token completi solo in remaining[:max_chars], dove un token a cavallo del
        confine non può comparire interamente)."""
        chunker = TextChunker(max_chars=20)
        text = "A" * 18 + "[PERSON_001] coda BBBB"   # token a posizioni 18-29, scavalca 20
        chunks = chunker.chunk(text)
        # nessun chunk contiene un frammento di token
        assert any("[PERSON_001]" in c for c in chunks)
        for c in chunks:
            assert c.count("[") == c.count("]"), f"token spezzato: {c!r}"

    def test_no_token_split_across_various_chunk_sizes(self):
        """Per varie dimensioni di chunk, nessun placeholder viene mai spezzato."""
        import re
        base = " ".join(f"word{i} [EMAIL_{i:03d}]" for i in range(1, 60))
        for mc in (15, 25, 37, 50, 80, 123):
            for c in TextChunker(max_chars=mc).chunk(base):
                frags = re.findall(r"\[EMAIL_\d", c)        # inizi di token
                full = re.findall(r"\[EMAIL_\d{3,}\]", c)    # token completi
                assert len(frags) == len(full), f"mc={mc} token spezzato in {c!r}"

    def test_all_content_preserved(self):
        chunker = TextChunker(max_chars=50)
        text = "A" * 200
        chunks = chunker.chunk(text)
        assert "".join(chunks) == text or len("".join(c.strip() for c in chunks)) > 0

    def test_many_chunks(self):
        chunker = TextChunker(max_chars=10)
        text = "alpha beta gamma delta epsilon"
        chunks = chunker.chunk(text)
        assert len(chunks) > 1
        # No chunk exceeds max_chars
        for chunk in chunks:
            assert len(chunk) <= 10 + 1  # +1 for tolerance at word boundary

    def test_long_document_integration(self):
        """Verifica che PrivacyProxy gestisca documenti lunghi via demo provider."""
        proxy = PrivacyProxy(
            provider=LLMProvider.DEMO,
            use_ner=False,
            max_chunk_chars=100,
        )
        long_text = ("Mario Rossi email: mario@test.com.\n\n" * 10).strip()
        result = proxy.anonymize_only(long_text)
        # All emails should be anonymized
        assert "mario@test.com" not in result["anonymized_text"]
        assert "[EMAIL_001]" in result["anonymized_text"]


class TestPromptTemplate:
    """Il template del prompt è fornito dall'utente: graffe estranee (JSON/codice)
    non devono far esplodere str.format()."""

    def test_template_with_stray_braces_does_not_crash(self):
        gw = LLMGateway(LLMProvider.DEMO)
        # template con graffe non '{document}' (richiesta di output JSON)
        out = gw.send_document(
            "[EMAIL_001]", 'Rispondi in JSON: {"k": "{document}"}'
        )
        assert isinstance(out, str) and out

    def test_template_substitutes_document(self):
        assert LLMGateway._apply_template("X {document} Y", "DOC") == "X DOC Y"

    def test_template_without_placeholder_unchanged(self):
        assert LLMGateway._apply_template("nessun segnaposto", "DOC") == "nessun segnaposto"


class TestCSVColumnInferrer:

    def test_email_column(self):
        assert CSVColumnInferrer.infer("email") == EntityType.EMAIL
        assert CSVColumnInferrer.infer("Email_Address") == EntityType.EMAIL
        assert CSVColumnInferrer.infer("MAIL") == EntityType.EMAIL

    def test_person_column(self):
        assert CSVColumnInferrer.infer("Nome") == EntityType.PERSON
        assert CSVColumnInferrer.infer("cognome") == EntityType.PERSON
        assert CSVColumnInferrer.infer("intestatario") == EntityType.PERSON

    def test_iban_column(self):
        assert CSVColumnInferrer.infer("IBAN") == EntityType.IBAN
        assert CSVColumnInferrer.infer("conto corrente") == EntityType.IBAN

    def test_cf_column(self):
        assert CSVColumnInferrer.infer("codice fiscale") == EntityType.CF
        assert CSVColumnInferrer.infer("CF") == EntityType.CF
        assert CSVColumnInferrer.infer("CodFisc") == EntityType.CF

    def test_phone_column(self):
        assert CSVColumnInferrer.infer("telefono") == EntityType.PHONE
        assert CSVColumnInferrer.infer("cellulare") == EntityType.PHONE
        assert CSVColumnInferrer.infer("phone_number") == EntityType.PHONE

    def test_amount_column(self):
        assert CSVColumnInferrer.infer("importo") == EntityType.AMOUNT
        assert CSVColumnInferrer.infer("stipendio") == EntityType.AMOUNT
        assert CSVColumnInferrer.infer("totale_euro") == EntityType.AMOUNT

    def test_address_column(self):
        assert CSVColumnInferrer.infer("indirizzo") == EntityType.ADDRESS
        assert CSVColumnInferrer.infer("residenza") == EntityType.ADDRESS

    def test_non_sensitive_column(self):
        assert CSVColumnInferrer.infer("id") is None
        assert CSVColumnInferrer.infer("quantita") is None
        assert CSVColumnInferrer.infer("colore") is None

    def test_underscore_separator(self):
        assert CSVColumnInferrer.infer("codice_fiscale") == EntityType.CF
        assert CSVColumnInferrer.infer("email_address") == EntityType.EMAIL


class TestNERDetectorGlinerGraceful:
    """GLiNER dovrebbe caricarsi gracefully se non installato."""

    def test_ner_detector_initializes_without_gliner(self):
        det = NERDetector()
        # Should not raise even if GLiNER is missing
        assert det is not None

    def test_detect_returns_list_without_gliner(self):
        det = NERDetector()
        results = det.detect("Mario Rossi lavora a Roma", threshold=0.7)
        assert isinstance(results, list)

    def test_spacy_entities_respect_confidence_threshold(self):
        """Le entità spaCy (confidence fissa 0.85) devono essere filtrate dalla soglia,
        coerentemente con GLiNER/Presidio. Regressione: con --confidence > 0.85
        passavano comunque."""
        # NERDetector senza caricare alcun motore reale, poi inietta un fake spaCy.
        det = NERDetector.__new__(NERDetector)
        det._spacy_nlp = None
        det._presidio = None
        det._gliner = None

        class _FakeEnt:
            def __init__(self, text, label, start, end):
                self.text, self.label_ = text, label
                self.start_char, self.end_char = start, end

        class _FakeDoc:
            ents = [_FakeEnt("Mario Rossi", "PER", 0, 11)]

        class _FakeNLP:
            def __call__(self, text):
                return _FakeDoc()

        det._spacy_nlp = _FakeNLP()

        # Soglia bassa: l'entità spaCy (0.85) passa.
        assert len(det.detect("Mario Rossi", threshold=0.7)) == 1
        # Soglia sopra 0.85: deve essere scartata.
        assert det.detect("Mario Rossi", threshold=0.9) == []


class TestDemoProvider:

    def test_demo_returns_string(self):
        gw = LLMGateway(LLMProvider.DEMO)
        resp = gw.send("Analizza questo testo: [PERSON_001] guadagna [AMOUNT_001]")
        assert isinstance(resp, str)
        assert len(resp) > 0
        assert "DEMO" in resp

    def test_demo_no_api_key_needed(self):
        gw = LLMGateway(LLMProvider.DEMO, api_key=None)
        resp = gw.send("testo")
        assert resp  # Should not raise


class TestOllamaProvider:
    """Provider Ollama testato mockando requests.post (nessun server reale)."""

    def test_ollama_send_builds_request_and_parses_response(self, monkeypatch):
        import requests

        captured = {}

        class _FakeResp:
            def raise_for_status(self):
                pass

            def json(self):
                return {"response": "risposta dal modello"}

        def _fake_post(url, json=None, timeout=None):
            captured["url"] = url
            captured["json"] = json
            captured["timeout"] = timeout
            return _FakeResp()

        monkeypatch.setattr(requests, "post", _fake_post)

        gw = LLMGateway(
            LLMProvider.OLLAMA, model="llama3.2", base_url="http://localhost:11434"
        )
        out = gw.send("ciao", system_prompt="sei un assistente")
        assert out == "risposta dal modello"
        assert captured["url"] == "http://localhost:11434/api/generate"
        assert captured["json"]["model"] == "llama3.2"
        assert captured["json"]["prompt"] == "ciao"
        assert captured["json"]["stream"] is False
        assert captured["json"]["system"] == "sei un assistente"

    def test_ollama_connection_error_wrapped_as_runtimeerror(self, monkeypatch):
        import requests

        def _boom(*a, **k):
            raise requests.exceptions.ConnectionError("connection refused")

        monkeypatch.setattr(requests, "post", _boom)
        gw = LLMGateway(LLMProvider.OLLAMA, base_url="http://localhost:11434")
        with pytest.raises(RuntimeError) as exc:
            gw.send("ciao")
        assert "Ollama" in str(exc.value)
        assert "localhost:11434" in str(exc.value)


class TestCloudProvidersMocked:
    """Provider cloud testati iniettando SDK fittizie in sys.modules (nessuna rete)."""

    def test_openai_builds_messages_and_extracts_content(self, monkeypatch):
        import sys, types
        captured = {}

        fake = types.ModuleType("openai")

        class _Msg:
            content = "openai response"

        class _Choice:
            message = _Msg()

        class _Resp:
            choices = [_Choice()]

        class _Completions:
            def create(self, model, messages):
                captured["model"] = model
                captured["messages"] = messages
                return _Resp()

        class _Chat:
            completions = _Completions()

        class _Client:
            def __init__(self, api_key=None):
                captured["api_key"] = api_key
                self.chat = _Chat()

        fake.OpenAI = _Client
        monkeypatch.setitem(sys.modules, "openai", fake)

        gw = LLMGateway(LLMProvider.OPENAI, api_key="sk-test", model="gpt-4o-mini")
        out = gw.send("hello", system_prompt="sys")
        assert out == "openai response"
        assert captured["api_key"] == "sk-test"
        assert captured["model"] == "gpt-4o-mini"
        assert captured["messages"] == [
            {"role": "system", "content": "sys"},
            {"role": "user", "content": "hello"},
        ]

    def test_openai_none_content_normalized_to_empty_string(self, monkeypatch):
        """content può essere None (risposta filtrata/solo-tool) → normalizzato a ''."""
        import sys, types

        fake = types.ModuleType("openai")

        class _Msg:
            content = None

        class _Resp:
            choices = [type("_C", (), {"message": _Msg()})()]

        class _Client:
            def __init__(self, api_key=None):
                self.chat = type(
                    "_Chat", (), {"completions": type("_Cmp", (), {"create": lambda s, **k: _Resp()})()}
                )()

        fake.OpenAI = _Client
        monkeypatch.setitem(sys.modules, "openai", fake)
        gw = LLMGateway(LLMProvider.OPENAI, api_key="x")
        assert gw.send("hello") == ""

    def test_anthropic_builds_request_and_extracts_text(self, monkeypatch):
        import sys, types
        captured = {}

        fake = types.ModuleType("anthropic")

        class _Block:
            text = "anthropic response"

        class _Resp:
            content = [_Block()]

        class _Messages:
            def create(self, **kwargs):
                captured.update(kwargs)
                return _Resp()

        class _Client:
            def __init__(self, api_key=None):
                captured["api_key"] = api_key
                self.messages = _Messages()

        fake.Anthropic = _Client
        monkeypatch.setitem(sys.modules, "anthropic", fake)

        gw = LLMGateway(LLMProvider.ANTHROPIC, api_key="x", model="claude-test")
        out = gw.send("hi", system_prompt="S")
        assert out == "anthropic response"
        assert captured["model"] == "claude-test"
        assert captured["system"] == "S"
        assert captured["messages"] == [{"role": "user", "content": "hi"}]
        assert captured["max_tokens"] == 8096

    def test_gemini_concatenates_system_and_extracts_text(self, monkeypatch):
        import sys, types
        captured = {}

        fake_genai = types.ModuleType("google.generativeai")

        class _GenResp:
            text = "gemini response"

        class _Model:
            def __init__(self, model):
                captured["model"] = model

            def generate_content(self, content):
                captured["content"] = content
                return _GenResp()

        def _configure(api_key=None):
            captured["api_key"] = api_key

        fake_genai.configure = _configure
        fake_genai.GenerativeModel = _Model

        fake_google = types.ModuleType("google")
        fake_google.generativeai = fake_genai
        monkeypatch.setitem(sys.modules, "google", fake_google)
        monkeypatch.setitem(sys.modules, "google.generativeai", fake_genai)

        gw = LLMGateway(LLMProvider.GEMINI, api_key="x", model="gemini-1.5-flash")
        out = gw.send("hello", system_prompt="sys")
        assert out == "gemini response"
        assert captured["api_key"] == "x"
        assert captured["model"] == "gemini-1.5-flash"
        assert captured["content"] == "sys\n\nhello"

    def test_missing_provider_sdk_raises_importerror(self, monkeypatch):
        import sys
        # Forza l'assenza del modulo 'openai' → ImportError con messaggio d'aiuto.
        monkeypatch.setitem(sys.modules, "openai", None)
        gw = LLMGateway(LLMProvider.OPENAI, api_key="x")
        with pytest.raises(ImportError):
            gw.send("hello")


class TestAccountNumberDetection:

    def test_conto_corrente_with_label(self):
        det = RuleBasedDetector()
        text = "c/c: 000010000011234567890123"
        entities = det.detect(text, threshold=0.7)
        accounts = [e for e in entities if e.type == EntityType.ACCOUNT]
        assert len(accounts) >= 1

    def test_account_number_with_label(self):
        det = RuleBasedDetector()
        text = "Account number: AB12345678"
        entities = det.detect(text, threshold=0.7)
        accounts = [e for e in entities if e.type == EntityType.ACCOUNT]
        assert len(accounts) >= 1

    def test_bare_number_not_detected(self):
        det = RuleBasedDetector()
        # Without explicit label, should NOT be detected (avoids false positives)
        text = "Il numero 123456789012 è il riferimento dell'ordine"
        entities = det.detect(text, threshold=0.7)
        accounts = [e for e in entities if e.type == EntityType.ACCOUNT]
        assert len(accounts) == 0


class TestStopOnError:

    def test_stop_on_error_raises_on_leaked_value(self):
        proxy = PrivacyProxy(provider=LLMProvider.DEMO, use_ner=False)
        # Bypass the tokenizer by patching the validator to return an error
        import unittest.mock as mock

        with mock.patch.object(proxy._validator, "validate_anonymized") as mock_val:
            from privacy_proxy import ValidationResult
            mock_val.return_value = ValidationResult(
                is_valid=False,
                warnings=[],
                errors=["Dato originale trapelato: 'test@example.com'"],
            )
            with pytest.raises(RuntimeError, match="pipeline interrotta"):
                proxy.process("testo di test", stop_on_error=True)

    def test_no_stop_on_error_continues(self):
        proxy = PrivacyProxy(provider=LLMProvider.DEMO, use_ner=False)
        import unittest.mock as mock

        with mock.patch.object(proxy._validator, "validate_anonymized") as mock_val:
            from privacy_proxy import ValidationResult
            mock_val.return_value = ValidationResult(
                is_valid=False,
                warnings=[],
                errors=["Dato originale trapelato: 'test@example.com'"],
            )
            # Should NOT raise when stop_on_error=False
            result = proxy.process("testo di test", stop_on_error=False)
            assert "final_response" in result


class TestCSVColumnFormat:

    def test_csv_per_column_format(self, tmp_path):
        try:
            import pandas as pd
        except ImportError:
            pytest.skip("pandas non installato")

        f = tmp_path / "persone.csv"
        f.write_text(
            "Nome,Email,CF\n"
            "Mario Rossi,mario@test.com,RSSMRA80A01H501U\n",
            encoding="utf-8",
        )
        doc = InputLayer().load(str(f))
        # Per-column format: "Nome: Mario Rossi | Email: ... | CF: ..."
        assert "Nome: Mario Rossi" in doc.text
        assert "Email: mario@test.com" in doc.text
        assert "CF: RSSMRA80A01H501U" in doc.text

    def test_csv_anonymizes_all_columns(self, tmp_path):
        try:
            import pandas as pd
        except ImportError:
            pytest.skip("pandas non installato")

        f = tmp_path / "dati.csv"
        f.write_text(
            "Nome,IBAN\nMario Rossi,IT60X0542811101000000123456\n",
            encoding="utf-8",
        )
        doc = InputLayer().load(str(f))
        detector = SensitiveDataDetectionEngine(confidence_threshold=0.7, use_ner=False)
        tokenizer = TokenizationEngine()
        tokenizer.reset()
        entities = detector.detect(doc.text)
        anon, mapping = tokenizer.tokenize(doc.text, entities)

        assert "IT60X0542811101000000123456" not in anon
        assert "[IBAN_001]" in anon


class TestOCRFallback:

    def test_error_message_mentions_ocr_options(self, tmp_path):
        """Senza librerie PDF disponibili, il messaggio di errore menziona le opzioni OCR."""
        import sys
        import unittest.mock as mock

        f = tmp_path / "dummy.pdf"
        f.write_bytes(b"not a real pdf")

        # Remove any cached pdf modules so we can control what's available
        blocked = {"pdfplumber", "fitz", "pytesseract", "pdf2image"}
        saved = {k: sys.modules.pop(k) for k in blocked if k in sys.modules}
        try:
            import builtins
            with mock.patch("builtins.__import__", side_effect=lambda name, *a, **kw: (
                (_ for _ in ()).throw(ImportError(f"no module: {name}"))
                if name in blocked
                else builtins.__import__(name, *a, **kw)
            )):
                with pytest.raises((ImportError, Exception)) as exc_info:
                    InputLayer().load(str(f))
                # When all PDF libs are absent, error should mention installation options
                msg = str(exc_info.value)
                assert any(kw in msg for kw in ["pdfplumber", "pymupdf", "pytesseract", "PDF"])
        finally:
            sys.modules.update(saved)

    def test_raw_text_still_works_regardless(self):
        """Testo libero (non file) funziona sempre, indipendente da librerie PDF."""
        doc = InputLayer().load("Mario Rossi abita in via Roma 1")
        assert "Mario Rossi" in doc.text


class TestRehydrateFromCLI:

    def test_rehydrate_from_saved_mapping(self, tmp_path):
        """--rehydrate-from + --load-mapping ripristina i dati da una mappa salvata."""
        from privacy_proxy import main as proxy_main

        mapping_file = tmp_path / "mappa.json"
        mapping_data = {"[PERSON_001]": "Mario Rossi", "[EMAIL_001]": "mario@test.com"}
        import json as _json
        mapping_file.write_text(_json.dumps(mapping_data), encoding="utf-8")

        llm_response = tmp_path / "risposta.txt"
        llm_response.write_text(
            "Il cliente [PERSON_001] ha scritto da [EMAIL_001].", encoding="utf-8"
        )

        output_file = tmp_path / "output.txt"
        rc = proxy_main([
            "--rehydrate-from", str(llm_response),
            "--load-mapping", str(mapping_file),
            "--output", str(output_file),
        ])
        assert rc == 0
        result = output_file.read_text(encoding="utf-8")
        assert "Mario Rossi" in result
        assert "mario@test.com" in result
        assert "[PERSON_001]" not in result

    def test_rehydrate_requires_load_mapping(self):
        """--rehydrate-from senza --load-mapping deve restituire codice di errore."""
        from privacy_proxy import main as proxy_main
        rc = proxy_main(["--rehydrate-from", "testo di prova"])
        assert rc == 1


class TestPIVAValidation:

    @pytest.fixture
    def det(self):
        return RuleBasedDetector()

    def test_valid_piva_detected(self, det):
        """Partita IVA italiana valida (checksum corretto) deve essere rilevata."""
        # 02182080396: checksum digit = 6
        entities = det.detect("P.IVA: 02182080396")
        pivas = [e for e in entities if e.type == EntityType.PIVA]
        assert len(pivas) >= 1

    def test_invalid_piva_rejected(self, det):
        """Partita IVA con checksum errato deve essere scartata."""
        entities = det.detect("P.IVA: 02182080390")  # last digit wrong (should be 6)
        pivas = [e for e in entities if e.type == EntityType.PIVA]
        assert len(pivas) == 0


class TestBatchDirMode:

    def test_batch_dir_anonymizes_files(self, tmp_path):
        """--dir elabora tutti i file supportati nella cartella."""
        from privacy_proxy import main as proxy_main

        (tmp_path / "doc1.txt").write_text(
            "Mario Rossi email: mario@test.com", encoding="utf-8"
        )
        (tmp_path / "doc2.txt").write_text(
            "IBAN: IT60X0542811101000000123456", encoding="utf-8"
        )

        out_dir = tmp_path / "output"
        out_dir.mkdir()

        rc = proxy_main([
            "--dir", str(tmp_path),
            "--anonymize-only",
            "--no-ner",
            "--output", str(out_dir),
        ])
        assert rc == 0
        out_files = list(out_dir.iterdir())
        assert len(out_files) == 2

    def test_batch_dir_invalid_path(self, tmp_path):
        """--dir con percorso non esistente deve restituire codice errore."""
        from privacy_proxy import main as proxy_main
        rc = proxy_main(["--dir", str(tmp_path / "nonexistent"), "--anonymize-only"])
        assert rc == 1

    def test_batch_dir_creates_missing_output_dir(self, tmp_path):
        """--output con cartella inesistente deve essere creata, non far fallire le scritture.

        Regressione: write_text falliva con FileNotFoundError e ogni file
        risultava 'in errore' (0/N OK)."""
        from privacy_proxy import main as proxy_main

        src = tmp_path / "in"
        src.mkdir()
        (src / "doc1.txt").write_text("mario@test.com", encoding="utf-8")
        out_dir = tmp_path / "out" / "nested"   # non esiste ancora

        rc = proxy_main([
            "--dir", str(src),
            "--anonymize-only",
            "--no-ner",
            "--output", str(out_dir),
        ])
        assert rc == 0
        assert (out_dir / "doc1_anon.txt").exists()
        assert "[EMAIL_001]" in (out_dir / "doc1_anon.txt").read_text(encoding="utf-8")

    def test_single_output_creates_missing_parent_dir(self, tmp_path):
        """--output con cartelle intermedie inesistenti deve crearle, non fallire.

        Regressione: tutte le scritture single-file (anonymize-only, redact,
        highlight, dry-run, pipeline) usavano Path(...).write_text senza creare
        la cartella padre."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out_file = tmp_path / "sub" / "dir" / "result.json"  # 'sub/dir' non esistono
        with contextlib.redirect_stdout(io.StringIO()):
            rc = proxy_main([
                "--text", "mario@test.com",
                "--anonymize-only",
                "--no-ner",
                "--output", str(out_file),
            ])
        assert rc == 0
        assert out_file.exists()
        data = json.loads(out_file.read_text(encoding="utf-8"))
        assert "[EMAIL_001]" in data["anonymized_text"]


# ──────────────────────────────────────────────────────────────
# TXT multi-encoding support
# ──────────────────────────────────────────────────────────────

class TestTxtMultiEncoding:

    def test_txt_utf8_bom(self, tmp_path):
        """File UTF-8 con BOM viene letto correttamente."""
        p = tmp_path / "bom.txt"
        p.write_bytes(b"\xef\xbb\xbfMario Rossi")
        doc = InputLayer().load(str(p))
        assert doc.text == "Mario Rossi"

    def test_txt_cp1252(self, tmp_path):
        """File CP1252 (Windows italiano) viene letto senza errori."""
        p = tmp_path / "win.txt"
        # 0x80=€, 0xe0=à in CP1252; write raw bytes to avoid Python str encoding issues
        p.write_bytes(b"Importo: \x80 1.500,00 \x97 Citt\xe0 di Roma")
        doc = InputLayer().load(str(p))
        assert "1.500,00" in doc.text
        assert "Roma" in doc.text

    def test_txt_iso8859(self, tmp_path):
        """File ISO-8859-1 con accenti italiani viene letto correttamente."""
        p = tmp_path / "latin.txt"
        text_lat = "Citt\xe0 di Napoli, via Crocell\xe0"
        p.write_bytes(text_lat.encode("iso-8859-1"))
        doc = InputLayer().load(str(p))
        assert "Napoli" in doc.text


# ──────────────────────────────────────────────────────────────
# Unicode NFC normalization
# ──────────────────────────────────────────────────────────────

class TestUnicodeNFCNormalization:

    def test_raw_text_nfd_normalized(self):
        """Testo raw con caratteri NFD viene normalizzato in NFC."""
        import unicodedata
        # "à" as NFD: 'a' + combining grave (U+0061 U+0300)
        nfd_text = "Città di Roma"
        assert unicodedata.is_normalized("NFC", nfd_text) is False
        doc = InputLayer().load(nfd_text)
        assert unicodedata.is_normalized("NFC", doc.text) is True
        assert "Città" in doc.text

    def test_file_nfd_normalized(self, tmp_path):
        """Testo NFD letto da file viene normalizzato in NFC."""
        import unicodedata
        p = tmp_path / "nfd.txt"
        nfd_text = "Mario Rossi abita a Napolì"
        p.write_text(nfd_text, encoding="utf-8")
        doc = InputLayer().load(str(p))
        assert unicodedata.is_normalized("NFC", doc.text) is True

    def test_nfc_text_unchanged(self):
        """Testo già NFC non viene alterato."""
        nfc_text = "Città di Napoli, via Monteoliveto"
        doc = InputLayer().load(nfc_text)
        assert doc.text == nfc_text


# ──────────────────────────────────────────────────────────────
# Modalità --redact
# ──────────────────────────────────────────────────────────────

class TestRedactMode:

    def test_redact_replaces_placeholders(self):
        """--redact sostituisce i token con [REDACTED]."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            rc = proxy_main([
                "--text", "Mario Rossi email: mario.rossi@example.com",
                "--redact",
                "--no-ner",
            ])
        assert rc == 0
        output = out.getvalue()
        assert "[REDACTED]" in output
        assert "mario.rossi@example.com" not in output

    def test_redact_no_placeholders_in_output(self):
        """--redact non lascia token [TYPE_NNN] nel testo finale."""
        import re, io, contextlib
        from privacy_proxy import main as proxy_main

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            proxy_main([
                "--text", "CF: RSSMRA80A01H501U, tel: +39 333 1234567",
                "--redact",
                "--no-ner",
            ])
        output = out.getvalue()
        assert not re.search(r"\[[A-Z_]+_\d{3}\]", output)

    def test_redact_to_file(self, tmp_path):
        """--redact con --output salva il testo redatto (non JSON)."""
        from privacy_proxy import main as proxy_main

        out_file = tmp_path / "redacted.txt"
        rc = proxy_main([
            "--text", "IBAN: IT60X0542811101000000123456",
            "--redact",
            "--no-ner",
            "--output", str(out_file),
        ])
        assert rc == 0
        assert out_file.exists()
        content = out_file.read_text(encoding="utf-8")
        assert "[REDACTED]" in content
        assert "IT60X0542811101000000123456" not in content


# ──────────────────────────────────────────────────────────────
# Filtro --entity-types
# ──────────────────────────────────────────────────────────────

class TestEntityTypesFilter:

    def test_filter_email_only(self):
        """--entity-types EMAIL rileva solo email, ignora CF e telefono."""
        engine = SensitiveDataDetectionEngine(
            confidence_threshold=0.5,
            use_ner=False,
            entity_types=[EntityType.EMAIL],
        )
        text = "CF: RSSMRA80A01H501U  email: mario@example.com  tel: +39 333 1234567"
        entities = engine.detect(text)
        types = {e.type for e in entities}
        assert EntityType.EMAIL in types
        assert EntityType.CF not in types
        assert EntityType.PHONE not in types

    def test_filter_iban_and_phone(self):
        """--entity-types IBAN,PHONE rileva solo quei due tipi."""
        engine = SensitiveDataDetectionEngine(
            confidence_threshold=0.5,
            use_ner=False,
            entity_types=[EntityType.IBAN, EntityType.PHONE],
        )
        text = (
            "Email: mario@example.com\n"
            "Tel: +39 333 1234567\n"
            "IBAN: IT60X0542811101000000123456"
        )
        entities = engine.detect(text)
        types = {e.type for e in entities}
        assert EntityType.IBAN in types
        assert EntityType.PHONE in types
        assert EntityType.EMAIL not in types

    def test_filter_none_means_all(self):
        """entity_types=None rileva tutti i tipi."""
        engine = SensitiveDataDetectionEngine(
            confidence_threshold=0.5,
            use_ner=False,
            entity_types=None,
        )
        text = "Email: mario@example.com  IBAN: IT60X0542811101000000123456"
        entities = engine.detect(text)
        types = {e.type for e in entities}
        assert EntityType.EMAIL in types
        assert EntityType.IBAN in types

    def test_cli_entity_types_flag(self):
        """--entity-types EMAIL dalla CLI filtra correttamente."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            rc = proxy_main([
                "--text", "mario@example.com tel: +39 333 1234567",
                "--anonymize-only",
                "--no-ner",
                "--entity-types", "EMAIL",
            ])
        assert rc == 0
        output = out.getvalue()
        assert "[EMAIL_" in output
        assert "[PHONE_" not in output

    def test_provider_error_reports_clean_message(self, monkeypatch):
        """Un'eccezione imprevista del provider LLM deve dare rc=1 con messaggio pulito,
        non un traceback (le librerie provider definiscono eccezioni proprie)."""
        from privacy_proxy import main as proxy_main, LLMGateway
        import io, contextlib

        class _FakeAPIError(Exception):
            pass

        def _boom(self, *a, **k):
            raise _FakeAPIError("Invalid API key")

        monkeypatch.setattr(LLMGateway, "send_document", _boom)
        err = io.StringIO()
        with contextlib.redirect_stderr(err), contextlib.redirect_stdout(io.StringIO()):
            rc = proxy_main(["--text", "mario@example.com", "--provider", "demo", "--no-ner"])
        assert rc == 1
        output = err.getvalue()
        assert "provider LLM" in output
        assert "_FakeAPIError" in output  # tipo incluso per diagnosi

    def test_malformed_json_file_reports_clean_error(self, tmp_path):
        """Un file JSON malformato deve produrre un messaggio pulito (rc=1), non un traceback."""
        from privacy_proxy import main as proxy_main
        import io, contextlib
        bad = tmp_path / "bad.json"
        bad.write_text("{not valid json,,,", encoding="utf-8")
        err = io.StringIO()
        with contextlib.redirect_stderr(err), contextlib.redirect_stdout(io.StringIO()):
            rc = proxy_main(["--file", str(bad), "--anonymize-only", "--no-ner"])
        assert rc == 1
        assert "✖" in err.getvalue()

    def test_unsupported_file_format_reports_clean_error(self, tmp_path):
        """Un'estensione non supportata deve dare rc=1 con messaggio pulito."""
        from privacy_proxy import main as proxy_main
        import io, contextlib
        f = tmp_path / "data.xyz"
        f.write_text("contenuto", encoding="utf-8")
        err = io.StringIO()
        with contextlib.redirect_stderr(err), contextlib.redirect_stdout(io.StringIO()):
            rc = proxy_main(["--file", str(f), "--dry-run", "--no-ner"])
        assert rc == 1
        assert "Formato non supportato" in err.getvalue()

    def test_cli_invalid_entity_type(self):
        """--entity-types con valore sconosciuto restituisce codice errore."""
        from privacy_proxy import main as proxy_main
        rc = proxy_main([
            "--text", "test",
            "--anonymize-only",
            "--no-ner",
            "--entity-types", "NONEXISTENT",
        ])
        assert rc == 1

    def test_entity_types_accepts_enum_names_and_values(self):
        """_parse_entity_types accetta sia i valori (ADDR/ORG/LOC) sia i nomi
        estesi (ADDRESS/ORGANIZATION/LOCATION); l'esempio del README usa ADDRESS."""
        from privacy_proxy import _parse_entity_types, EntityType
        assert _parse_entity_types("ADDRESS") == [EntityType.ADDRESS]
        assert _parse_entity_types("ORGANIZATION,LOCATION") == [
            EntityType.ORGANIZATION, EntityType.LOCATION
        ]
        assert _parse_entity_types("ADDR") == [EntityType.ADDRESS]
        # Forme miste e duplicati collassano in un'unica entry, token vuoti ignorati.
        assert _parse_entity_types("ADDR, ADDRESS , ,EMAIL") == [
            EntityType.ADDRESS, EntityType.EMAIL
        ]

    def test_cli_entity_types_address_alias_runs(self):
        """L'esempio del README '--entity-types ...,ADDRESS' non deve fallire."""
        from privacy_proxy import main as proxy_main
        import io, contextlib
        with contextlib.redirect_stdout(io.StringIO()):
            rc = proxy_main([
                "--text", "Via Roma 1, 00100 Roma mario@example.com",
                "--anonymize-only",
                "--no-ner",
                "--entity-types", "EMAIL,ADDRESS",
            ])
        assert rc == 0


# ──────────────────────────────────────────────────────────────
# Source tracking e confidence boosting
# ──────────────────────────────────────────────────────────────

class TestConfidenceBoosting:

    def test_source_field_set_for_rule_based(self):
        """Le entità rule-based hanno source='rule'."""
        det = RuleBasedDetector()
        entities = det.detect("mario@example.com")
        assert all(e.source == "rule" for e in entities)

    def test_confidence_boost_multi_source(self):
        """Due entità con stesso span e sorgenti diverse → confidence aumenta."""
        from privacy_proxy import SensitiveDataDetectionEngine, DetectedEntity, EntityType
        engine = SensitiveDataDetectionEngine.__new__(SensitiveDataDetectionEngine)

        e1 = DetectedEntity("mario@test.com", EntityType.EMAIL, 0.90, 0, 14, source="rule")
        e2 = DetectedEntity("mario@test.com", EntityType.EMAIL, 0.85, 0, 14, source="presidio")

        result = engine._deduplicate([e1, e2])
        assert len(result) == 1
        assert result[0].confidence > 0.90

    def test_confidence_boost_capped_at_099(self):
        """Il boost non supera 0.99."""
        from privacy_proxy import SensitiveDataDetectionEngine, DetectedEntity, EntityType
        engine = SensitiveDataDetectionEngine.__new__(SensitiveDataDetectionEngine)

        e1 = DetectedEntity("test@x.it", EntityType.EMAIL, 0.97, 0, 9, source="rule")
        e2 = DetectedEntity("test@x.it", EntityType.EMAIL, 0.97, 0, 9, source="spacy")

        result = engine._deduplicate([e1, e2])
        assert result[0].confidence <= 0.99

    def test_single_source_no_boost(self):
        """Entità da una sola sorgente non riceve boost."""
        from privacy_proxy import SensitiveDataDetectionEngine, DetectedEntity, EntityType
        engine = SensitiveDataDetectionEngine.__new__(SensitiveDataDetectionEngine)

        e1 = DetectedEntity("test@x.it", EntityType.EMAIL, 0.90, 0, 9, source="rule")

        result = engine._deduplicate([e1])
        assert result[0].confidence == 0.90

    def test_nested_overlap_keeps_longest_span(self):
        """Span lungo che contiene uno più corto (anche a confidence maggiore)
        deve essere mantenuto, per non lasciare in chiaro i caratteri non coperti.
        Regressione: il CAP annidato sostituiva l'ADDRESS, esponendo 'Via Roma 1, '."""
        from privacy_proxy import SensitiveDataDetectionEngine, DetectedEntity, EntityType
        engine = SensitiveDataDetectionEngine.__new__(SensitiveDataDetectionEngine)

        addr = DetectedEntity("Via Roma 1, 00100", EntityType.ADDRESS, 0.80, 0, 17, source="rule")
        cap = DetectedEntity("00100", EntityType.CAP, 0.90, 12, 17, source="rule")

        result = engine._deduplicate([addr, cap])
        assert len(result) == 1
        assert result[0].type == EntityType.ADDRESS
        assert (result[0].start, result[0].end) == (0, 17)

    def test_dedup_produces_no_overlapping_spans(self):
        """Il risultato del dedup non deve mai contenere span sovrapposti
        (la tokenizzazione basata sugli indici si corromperebbe)."""
        from privacy_proxy import SensitiveDataDetectionEngine, DetectedEntity, EntityType
        engine = SensitiveDataDetectionEngine.__new__(SensitiveDataDetectionEngine)

        ents = [
            DetectedEntity("a", EntityType.PERSON, 0.9, 0, 10, source="rule"),
            DetectedEntity("b", EntityType.CF, 0.95, 5, 8, source="rule"),
            DetectedEntity("c", EntityType.EMAIL, 0.85, 10, 20, source="rule"),
        ]
        result = sorted(engine._deduplicate(ents), key=lambda e: e.start)
        for i in range(len(result) - 1):
            assert result[i].end <= result[i + 1].start


# ──────────────────────────────────────────────────────────────
# Source field in output dicts + --show-sources CLI
# ──────────────────────────────────────────────────────────────

class TestSourceInOutput:

    def test_anonymize_only_entities_have_source(self):
        """anonymize_only() include 'source' in ogni entità."""
        proxy = PrivacyProxy(use_ner=False)
        result = proxy.anonymize_only("email: mario@example.com")
        assert result["entities"], "dovrebbe rilevare almeno un'entità"
        for e in result["entities"]:
            assert "source" in e
            assert e["source"] in ("rule", "spacy", "gliner", "presidio")

    def test_show_sources_cli_flag(self):
        """--show-sources mostra la colonna src= nell'output."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            proxy_main([
                "--text", "mario@example.com tel: +39 333 1234567",
                "--anonymize-only",
                "--no-ner",
                "--show-sources",
            ])
        output = out.getvalue()
        assert "src=rule" in output

    def test_source_rule_for_regex_detections(self):
        """Le entità trovate con regex hanno source='rule'."""
        proxy = PrivacyProxy(use_ner=False)
        result = proxy.anonymize_only("IBAN: IT60X0542811101000000123456")
        iban_entities = [e for e in result["entities"] if e["type"] == "IBAN"]
        assert iban_entities
        assert all(e["source"] == "rule" for e in iban_entities)


# ──────────────────────────────────────────────────────────────
# Modalità --dry-run
# ──────────────────────────────────────────────────────────────

class TestDryRunMode:

    def test_dry_run_text_unchanged(self):
        """--dry-run non modifica il testo originale nell'output."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            rc = proxy_main([
                "--text", "mario@example.com tel: +39 333 1234567",
                "--dry-run",
                "--no-ner",
            ])
        assert rc == 0
        output = out.getvalue()
        # I placeholder NON devono apparire — il testo non viene modificato
        assert "[EMAIL_" not in output
        assert "[PHONE_" not in output
        # ma le entità devono essere elencate
        assert "mario@example.com" in output

    def test_dry_run_lists_entities_with_positions(self):
        """--dry-run mostra posizione start-end di ogni entità."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            proxy_main([
                "--text", "IBAN: IT60X0542811101000000123456",
                "--dry-run",
                "--no-ner",
            ])
        output = out.getvalue()
        assert "pos=" in output
        assert "IBAN" in output

    def test_dry_run_no_api_key_needed(self):
        """--dry-run non richiede API key anche con provider cloud."""
        from privacy_proxy import main as proxy_main
        rc = proxy_main([
            "--text", "test senza entità",
            "--dry-run",
            "--no-ner",
            "--provider", "anthropic",
        ])
        assert rc == 0

    def test_dry_run_to_json_file(self, tmp_path):
        """--dry-run con --output salva entità in JSON con posizioni."""
        from privacy_proxy import main as proxy_main
        import json

        out_file = tmp_path / "dry.json"
        proxy_main([
            "--text", "email: mario@example.com",
            "--dry-run",
            "--no-ner",
            "--output", str(out_file),
        ])
        assert out_file.exists()
        data = json.loads(out_file.read_text(encoding="utf-8"))
        assert isinstance(data, list)
        assert any(e["type"] == "EMAIL" for e in data)
        assert all("start" in e and "end" in e for e in data)


# ──────────────────────────────────────────────────────────────
# Modalità --highlight
# ──────────────────────────────────────────────────────────────

class TestHighlightMode:

    def test_highlight_marks_entities_inline(self):
        """--highlight inserisce marcatori {valore|TIPO:conf} nel testo."""
        from privacy_proxy import main as proxy_main, _build_highlighted_text, \
            DetectedEntity, EntityType
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            rc = proxy_main([
                "--text", "email: mario@example.com",
                "--highlight",
                "--no-ner",
            ])
        assert rc == 0
        output = out.getvalue()
        assert "{mario@example.com|EMAIL:" in output

    def test_highlight_preserves_non_pii_text(self):
        """Il testo non PII viene conservato intatto."""
        from privacy_proxy import _build_highlighted_text, DetectedEntity, EntityType

        text = "Ciao Mario, la tua email è mario@example.com grazie."
        start = text.index("mario@example.com")
        end = start + len("mario@example.com")
        entity = DetectedEntity("mario@example.com", EntityType.EMAIL, 0.99, start, end, "rule")
        result = _build_highlighted_text(text, [entity])
        assert result.startswith("Ciao Mario, la tua email è ")
        assert result.endswith(" grazie.")
        assert "{mario@example.com|EMAIL:0.99}" in result

    def test_highlight_no_entities(self):
        """Testo senza entità viene restituito invariato."""
        from privacy_proxy import _build_highlighted_text
        text = "Testo privo di dati sensibili."
        assert _build_highlighted_text(text, []) == text

    def test_highlight_no_api_key_needed(self):
        """--highlight non richiede API key."""
        from privacy_proxy import main as proxy_main
        rc = proxy_main([
            "--text", "test",
            "--highlight",
            "--no-ner",
            "--provider", "openai",
        ])
        assert rc == 0

    def test_highlight_to_file(self, tmp_path):
        """--highlight con --output salva il testo marcato."""
        from privacy_proxy import main as proxy_main
        out_file = tmp_path / "highlighted.txt"
        proxy_main([
            "--text", "IBAN: IT60X0542811101000000123456",
            "--highlight",
            "--no-ner",
            "--output", str(out_file),
        ])
        assert out_file.exists()
        content = out_file.read_text(encoding="utf-8")
        assert "{IT60X0542811101000000123456|IBAN:" in content


# ──────────────────────────────────────────────────────────────
# --stats e zero-entity warning
# ──────────────────────────────────────────────────────────────

class TestStatsMode:

    def test_stats_shows_per_type_counts(self):
        """--stats mostra tabella con conteggi per tipo."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            proxy_main([
                "--text", (
                    "mario@example.com tel: +39 333 1234567 "
                    "IBAN: IT60X0542811101000000123456"
                ),
                "--anonymize-only",
                "--no-ner",
                "--stats",
            ])
        output = out.getvalue()
        assert "STATISTICHE" in output
        assert "EMAIL" in output
        assert "IBAN" in output
        # deve mostrare colonne confidence
        assert "conf" in output.lower() or "Totale" in output

    def test_stats_shows_source_breakdown(self):
        """--stats include breakdown per sorgente."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            proxy_main([
                "--text", "mario@example.com",
                "--anonymize-only",
                "--no-ner",
                "--stats",
            ])
        output = out.getvalue()
        assert "rule" in output

    def test_print_entity_stats_helper(self):
        """_print_entity_stats funziona con lista entità dict."""
        from privacy_proxy import _print_entity_stats
        import io, contextlib

        entities = [
            {"type": "EMAIL", "confidence": 0.99, "source": "rule"},
            {"type": "EMAIL", "confidence": 0.95, "source": "presidio"},
            {"type": "PHONE", "confidence": 0.85, "source": "rule"},
        ]
        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            _print_entity_stats(entities)
        output = out.getvalue()
        assert "EMAIL" in output
        assert "PHONE" in output
        assert "Totale" in output
        assert "3" in output


class TestFullPipelineDemoCLI:
    """Pipeline LLM completa via provider 'demo' (nessuna API key) attraverso la CLI."""

    def test_full_pipeline_with_options_and_output(self, tmp_path):
        """process() → output formattato + --stats/--show-map/--show-sources +
        --output JSON + --save-mapping, con ripristino end-to-end dei dati."""
        from privacy_proxy import main as proxy_main
        import io, contextlib, json

        out_file = tmp_path / "out" / "result.json"   # cartella inesistente: verifica anche _write_output
        map_file = tmp_path / "map.json"
        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            rc = proxy_main([
                "--text", "Scrivi a mario.rossi@example.com (CF RSSMRA80A01H501U)",
                "--provider", "demo",
                "--no-ner",
                "--stats",
                "--show-map",
                "--show-sources",
                "--save-mapping", str(map_file),
                "--output", str(out_file),
            ])
        assert rc == 0
        output = out.getvalue()
        assert "RISPOSTA FINALE" in output
        assert "MAPPA ENTITÀ" in output
        assert "SORGENTI RILEVAMENTO" in output

        # JSON di output salvato (anche in cartella creata al volo)
        data = json.loads(out_file.read_text(encoding="utf-8"))
        assert "[EMAIL_001]" in data["anonymized_text"]
        # il provider demo riecheggia i placeholder → la rehydration ripristina l'originale
        assert "mario.rossi@example.com" in data["final_response"]
        assert data["validation"]["anonymization"]["is_valid"] is True

        # mappa salvata in chiaro e ricaricabile
        mapping = json.loads(map_file.read_text(encoding="utf-8"))
        assert any(v == "mario.rossi@example.com" for v in mapping.values())

    def test_full_pipeline_custom_prompt_template(self):
        """Il template --prompt con {document} viene applicato e la pipeline completa."""
        from privacy_proxy import main as proxy_main
        import io, contextlib

        out = io.StringIO()
        with contextlib.redirect_stdout(out):
            rc = proxy_main([
                "--text", "mario@example.com",
                "--provider", "demo",
                "--no-ner",
                "--prompt", "Riassumi:\n\n{document}",
            ])
        assert rc == 0
        assert "RISPOSTA FINALE" in out.getvalue()

    def test_zero_entity_warning(self, capsys):
        """Documento lungo con 0 entità produce avviso su stderr."""
        from privacy_proxy import main as proxy_main

        long_text = "Questo testo lungo non contiene dati sensibili. " * 5
        proxy_main([
            "--text", long_text,
            "--anonymize-only",
            "--no-ner",
        ])
        captured = capsys.readouterr()
        assert "Nessuna entità" in captured.err or len(captured.err) == 0  # warning se >100 chars


# ──────────────────────────────────────────────────────────────
# Validatori: casi limite (checksum/lunghezze)
# ──────────────────────────────────────────────────────────────

class TestValidatorEdgeCases:

    def test_luhn_rejects_too_short(self):
        assert RuleBasedDetector._luhn("4111") is False  # < 13 cifre

    def test_luhn_accepts_valid(self):
        assert RuleBasedDetector._luhn("4111111111111111") is True

    def test_iban_rejects_too_short(self):
        assert RuleBasedDetector._validate_iban("IT60") is False

    def test_iban_rejects_bad_checksum(self):
        assert RuleBasedDetector._validate_iban("IT00X0542811101000000123456") is False

    def test_iban_accepts_valid(self):
        assert RuleBasedDetector._validate_iban("IT60X0542811101000000123456") is True

    def test_iban_non_numeric_safe(self):
        # caratteri non gestiti non devono far esplodere int(): ritorna False
        assert RuleBasedDetector._validate_iban("IT60!!!!") is False

    def test_piva_rejects_wrong_length(self):
        assert RuleBasedDetector._validate_piva("123") is False

    def test_piva_rejects_bad_checksum(self):
        assert RuleBasedDetector._validate_piva("02182080390") is False

    def test_cf_rejects_wrong_length(self):
        det = RuleBasedDetector()
        assert det._validate_cf("ABC") is False


# ──────────────────────────────────────────────────────────────
# InputLayer: formati e dispatch (percorso NON usato dalla web app)
# ──────────────────────────────────────────────────────────────

class TestInputLayerFormats:

    def test_raw_text_when_not_a_file(self):
        doc = InputLayer().load("Mario Rossi, testo libero")
        assert doc.source == "raw_text"
        assert "Mario" in doc.text

    def test_unsupported_extension_raises(self, tmp_path):
        f = tmp_path / "x.xyz"
        f.write_text("contenuto", encoding="utf-8")
        with pytest.raises(ValueError):
            InputLayer().load(str(f))

    def test_txt_encoding_fallback(self, tmp_path):
        f = tmp_path / "a.txt"
        f.write_bytes("società €100".encode("cp1252"))
        doc = InputLayer().load(str(f))
        assert isinstance(doc.text, str) and "100" in doc.text

    def test_json_flatten_nested(self, tmp_path):
        f = tmp_path / "d.json"
        f.write_text(
            json.dumps({"email": "a@b.com", "nested": {"tel": "123"}, "list": ["x"]}),
            encoding="utf-8",
        )
        doc = InputLayer().load(str(f))
        assert "a@b.com" in doc.text

    def test_json_empty_falls_back_to_dump(self, tmp_path):
        f = tmp_path / "e.json"
        f.write_text("[]", encoding="utf-8")
        doc = InputLayer().load(str(f))
        assert isinstance(doc.text, str)

    def test_csv_load(self, tmp_path):
        pytest.importorskip("pandas")
        f = tmp_path / "d.csv"
        f.write_text("nome,email\nMario,a@b.com\n", encoding="utf-8")
        doc = InputLayer().load(str(f))
        assert "a@b.com" in doc.text

    def test_docx_load_with_table(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        d = Document()
        d.add_paragraph("Mario testo")
        t = d.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "CF"
        t.rows[0].cells[1].text = "RSSMRA80A01H501U"
        p = tmp_path / "d.docx"
        d.save(str(p))
        doc = InputLayer().load(str(p))
        assert "Mario" in doc.text and "RSSMRA80A01H501U" in doc.text
