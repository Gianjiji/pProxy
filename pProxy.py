#!/usr/bin/env python3
"""
Privacy Proxy per LLM v1.1
Anonimizza documenti prima di inviarli agli LLM cloud, poi ripristina i dati originali.

Architettura (da specifica):
  Modulo 1 - Input Layer          (PDF, CSV, TXT, DOCX, JSON → testo + inferenza colonne CSV)
  Modulo 2 - Detection Engine     (regex + NER ibrido: spaCy, GLiNER, Presidio)
  Modulo 3 - Tokenization Engine  (sostituzione deterministica con placeholder)
  Modulo 4 - Secure Mapping Store (mappa locale, AES-256 opzionale)
  Modulo 5 - LLM Gateway          (OpenAI, Anthropic, Gemini, Ollama, Demo)
  Modulo 6 - Rehydration Engine   (ripristino dati originali nella risposta)
  Modulo 7 - Validation Layer     (verifica token non sostituiti / inventati)
"""

from __future__ import annotations

import re
import json
import os
import sys
import logging
import argparse
import unicodedata
from pathlib import Path
from typing import Optional, Dict, List, Tuple, Any
from dataclasses import dataclass, field
from enum import Enum

# NB: la configurazione del logging (basicConfig) avviene in main(), non a livello
# di modulo: importare pProxy/privacy_proxy come libreria non deve riconfigurare
# il root logger dell'applicazione ospite (anti-pattern). Qui si ottiene solo il
# logger del modulo, lasciando all'applicazione il controllo degli handler.
logger = logging.getLogger(__name__)


def _configure_logging(verbose: bool = False) -> None:
    """Configura il logging per l'uso da riga di comando.

    Idempotente: basicConfig aggiunge un handler solo se il root non ne ha già,
    quindi invocazioni ripetute di main() (es. nei test) sono innocue.
    """
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        datefmt="%H:%M:%S",
    )
    if verbose:
        logging.getLogger().setLevel(logging.DEBUG)


# ──────────────────────────────────────────────────────────────
# DATA TYPES CONDIVISI
# ──────────────────────────────────────────────────────────────

class EntityType(str, Enum):
    PERSON       = "PERSON"
    ORGANIZATION = "ORG"
    LOCATION     = "LOC"
    ADDRESS      = "ADDR"
    EMAIL        = "EMAIL"
    PHONE        = "PHONE"
    IBAN         = "IBAN"
    CF           = "CF"          # Codice Fiscale
    PIVA         = "PIVA"        # Partita IVA
    CARD         = "CARD"        # Carta di credito
    DATE         = "DATE"
    AMOUNT       = "AMOUNT"
    CAP          = "CAP"
    ACCOUNT      = "ACCOUNT"
    # ── Identificativi documenti (IT + EN) ──
    PASSPORT     = "PASSPORT"      # Passaporto / passport
    ID_CARD      = "ID_CARD"       # Carta d'identità / national ID card
    DRIVER_LICENSE = "DRIVING_LICENSE"  # Patente / driving licence
    LICENSE_PLATE  = "PLATE"       # Targa / license plate
    TAX_ID       = "TAX_ID"        # SSN / NINO / EU VAT / generic national tax id
    # ── Dati finanziari aggiuntivi ──
    BIC          = "BIC"           # BIC / SWIFT
    CVV          = "CVV"           # Card security code
    CARD_EXPIRY  = "CARD_EXPIRY"   # Scadenza carta
    # ── Identificativi di rete ──
    IP_ADDRESS   = "IP"            # IPv4 / IPv6
    MAC_ADDRESS  = "MAC"           # MAC address
    URL          = "URL"           # URL / link
    USERNAME     = "USERNAME"      # @handle / username
    # ── Identificativi lavorativi e sanitari ──
    EMPLOYEE_ID  = "EMPLOYEE_ID"   # Matricola / employee id
    MEDICAL_ID   = "MEDICAL_ID"    # Cartella clinica / patient id
    INSURANCE_ID = "INSURANCE_ID"  # Assicurazione / policy number


@dataclass
class DetectedEntity:
    value: str
    type: EntityType
    confidence: float
    start: int
    end: int
    source: str = "rule"  # "rule", "spacy", "gliner", "presidio"


@dataclass
class RawTextDocument:
    text: str
    source: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ValidationResult:
    is_valid: bool
    warnings: List[str]
    errors: List[str]


# ──────────────────────────────────────────────────────────────
# MODULO 1 – INPUT LAYER
# ──────────────────────────────────────────────────────────────

class InputLayer:
    """Converte PDF, CSV, TXT, DOCX, JSON in RawTextDocument uniforme."""

    def load(self, source: str) -> RawTextDocument:
        path = Path(source)
        try:
            is_file = path.exists()
        except OSError:
            is_file = False
        if not is_file:
            doc = RawTextDocument(text=source, source="raw_text")
            doc.text = unicodedata.normalize("NFC", doc.text)
            return doc

        ext = path.suffix.lower()
        dispatch = {
            ".txt":  self._txt,
            ".pdf":  self._pdf,
            ".csv":  self._csv,
            ".docx": self._docx,
            ".json": self._json_file,
        }
        loader = dispatch.get(ext)
        if loader is None:
            raise ValueError(f"Formato non supportato: {ext}")
        doc = loader(path)
        doc.text = unicodedata.normalize("NFC", doc.text)
        return doc

    # ── loaders ──────────────────────────────────────────────

    @staticmethod
    def _txt(path: Path) -> RawTextDocument:
        text: Optional[str] = None
        for enc in ("utf-8-sig", "utf-8", "cp1252", "iso-8859-1"):
            try:
                text = path.read_text(encoding=enc)
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            text = path.read_text(encoding="utf-8", errors="replace")
        return RawTextDocument(text=text, source=str(path))

    @staticmethod
    def _pdf(path: Path) -> RawTextDocument:
        text = ""

        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text() or ""
                    text += extracted + "\n"
        except ImportError:
            pass

        if not text.strip():
            try:
                import fitz  # pymupdf
                doc = fitz.open(str(path))
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except ImportError:
                pass

        # OCR fallback for scanned PDFs (spec: "OCR se necessario")
        if not text.strip():
            try:
                import pytesseract
                from pdf2image import convert_from_path
                images = convert_from_path(str(path))
                for img in images:
                    text += pytesseract.image_to_string(img, lang="ita+eng") + "\n"
            except ImportError:
                pass

        if not text.strip():
            raise ImportError(
                "Nessuna libreria PDF disponibile. Installare:\n"
                "  pip install pdfplumber        # PDF testuali\n"
                "  pip install pymupdf           # PDF testuali (alternativa)\n"
                "  pip install pytesseract pdf2image  # PDF scansionati (OCR)"
            )
        return RawTextDocument(text=text, source=str(path), metadata={"type": "pdf"})

    @staticmethod
    def _csv(path: Path) -> RawTextDocument:
        """
        CSV: per ogni colonna inferisce il tipo (CSVColumnInferrer) e serializza in testo.
        Formato per riga: 'COLONNA: valore | COLONNA2: valore2'
        Le colonne classificate come sensibili vengono annotate per aiutare il detector.
        """
        try:
            import pandas as pd
            df = pd.read_csv(path)
            col_types = {col: CSVColumnInferrer.infer(col) for col in df.columns}
            sensitive_cols = {col: t for col, t in col_types.items() if t is not None}
            lines = []
            for _, row in df.iterrows():
                parts = [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
                lines.append(" | ".join(parts))
            text = "\n".join(lines)
            return RawTextDocument(
                text=text,
                source=str(path),
                metadata={
                    "type": "csv",
                    "columns": list(df.columns),
                    "rows": len(df),
                    "sensitive_columns": {c: t.value for c, t in sensitive_cols.items()},
                },
            )
        except ImportError:
            return RawTextDocument(
                text=path.read_text(encoding="utf-8", errors="replace"),
                source=str(path),
                metadata={"type": "csv"},
            )

    @staticmethod
    def _docx(path: Path) -> RawTextDocument:
        try:
            from docx import Document
            doc = Document(str(path))
            parts: List[str] = [p.text for p in doc.paragraphs if p.text.strip()]
            # Estrai anche il contenuto delle tabelle (celle spesso contengono dati sensibili)
            for table in doc.tables:
                for row in table.rows:
                    # Deduplicare celle adiacenti con lo stesso testo (celle unite)
                    seen_cell: Optional[str] = None
                    cell_values = []
                    for cell in row.cells:
                        val = cell.text.strip()
                        if val and val != seen_cell:
                            cell_values.append(val)
                            seen_cell = val
                    if cell_values:
                        parts.append(" | ".join(cell_values))
            text = "\n".join(parts)
            return RawTextDocument(text=text, source=str(path), metadata={"type": "docx"})
        except ImportError:
            raise ImportError("Installare python-docx: pip install python-docx")

    @staticmethod
    def _json_file(path: Path) -> RawTextDocument:
        data = json.loads(path.read_text(encoding="utf-8"))
        # Serializzazione annotata: "campo: valore" – stessa logica del CSV per-colonna,
        # così il rilevatore ha il nome del campo come contesto (es. "email: mario@test.com")
        lines = InputLayer._json_flatten(data)
        text = "\n".join(lines) if lines else json.dumps(data, ensure_ascii=False, indent=2)
        return RawTextDocument(
            text=text,
            source=str(path),
            metadata={"type": "json"},
        )

    @staticmethod
    def _json_flatten(obj: Any, prefix: str = "") -> List[str]:
        """Appiattisce un oggetto JSON in righe 'chiave: valore' per agevolare il detector."""
        lines: List[str] = []
        if isinstance(obj, dict):
            for key, value in obj.items():
                label = f"{prefix}.{key}" if prefix else key
                if isinstance(value, (dict, list)):
                    lines.extend(InputLayer._json_flatten(value, label))
                elif value is not None and str(value).strip():
                    lines.append(f"{key}: {value}")
        elif isinstance(obj, list):
            for item in obj:
                if isinstance(item, (dict, list)):
                    lines.extend(InputLayer._json_flatten(item, prefix))
                elif item is not None and str(item).strip():
                    lines.append(str(item))
        elif obj is not None:
            lines.append(str(obj))
        return lines


# ──────────────────────────────────────────────────────────────
# MODULO 1B – CSV COLUMN TYPE INFERRER
# ──────────────────────────────────────────────────────────────

class CSVColumnInferrer:
    """
    Inferisce il tipo di dato sensibile da un header di colonna CSV.
    Spec: "Per ogni colonna: inferire il tipo, classificare il dato".
    Restituisce EntityType o None se la colonna non è sensibile.
    """

    # Parole chiave nei nomi di colonna → tipo di entità
    _HEADER_MAP: List[Tuple[List[str], EntityType]] = [
        (["nome", "cognome", "name", "surname", "persona", "intestatario", "titolare",
          "nombre", "apellido", "nom", "prénom", "prenom", "sobrenome", "vorname",
          "nachname", "kunde"], EntityType.PERSON),
        (["email", "mail", "e-mail", "posta", "correo", "courriel", "correio"], EntityType.EMAIL),
        (["telefono", "tel", "phone", "cellulare", "mobile", "fax", "teléfono", "telefono",
          "téléphone", "telefone", "telefon", "móvil", "movil", "handy", "celular"], EntityType.PHONE),
        (["iban", "conto corrente", "conto", "account"], EntityType.IBAN),
        (["codice fiscale", "cf", "fiscal", "codfis", "codfisc"], EntityType.CF),
        (["partita iva", "p.iva", "piva", "vat"], EntityType.PIVA),
        (["indirizzo", "address", "via", "street", "residenza", "domicilio",
          "dirección", "direccion", "calle", "adresse", "rue", "morada", "rua",
          "anschrift", "straße", "strasse"], EntityType.ADDRESS),
        (["cap", "zip", "postal", "código postal", "codigo postal", "code postal",
          "plz", "postleitzahl"], EntityType.CAP),
        (["data", "date", "nascita", "birth", "nato", "fecha", "nacimiento",
          "naissance", "geburt", "geburtsdatum", "nascimento"], EntityType.DATE),
        (["importo", "amount", "totale", "saldo", "stipendio", "reddito", "costo", "prezzo", "euro"], EntityType.AMOUNT),
        (["carta", "card", "cc", "cvv", "credito"], EntityType.CARD),
        (["citta", "città", "city", "comune", "paese", "localita", "località", "location"], EntityType.LOCATION),
        (["organizzazione", "azienda", "societa", "società", "company", "org", "datore"], EntityType.ORGANIZATION),
        (["passaporto", "passport"], EntityType.PASSPORT),
        (["carta identita", "carta d'identita", "carta d identita", "documento", "id card", "identity"], EntityType.ID_CARD),
        (["patente", "driving licence", "driving license", "driver licence", "driver license"], EntityType.DRIVER_LICENSE),
        (["targa", "license plate", "number plate", "plate"], EntityType.LICENSE_PLATE),
        (["ssn", "social security", "national insurance", "nino", "tax id", "vat"], EntityType.TAX_ID),
        (["bic", "swift"], EntityType.BIC),
        (["cvv", "cvc", "security code"], EntityType.CVV),
        (["scadenza", "expiry", "exp date", "valid thru"], EntityType.CARD_EXPIRY),
        (["ip", "indirizzo ip", "ip address"], EntityType.IP_ADDRESS),
        (["mac", "mac address"], EntityType.MAC_ADDRESS),
        (["url", "link", "sito", "website"], EntityType.URL),
        (["username", "user", "utente", "handle", "nickname", "login"], EntityType.USERNAME),
        (["matricola", "employee id", "badge", "staff id", "emp"], EntityType.EMPLOYEE_ID),
        (["cartella clinica", "medical record", "patient id", "paziente", "health id"], EntityType.MEDICAL_ID),
        (["assicurazione", "polizza", "insurance", "policy"], EntityType.INSURANCE_ID),
    ]

    @classmethod
    def infer(cls, header: str) -> Optional[EntityType]:
        normalized = header.lower().strip().replace("_", " ").replace("-", " ")
        for keywords, etype in cls._HEADER_MAP:
            if any(kw in normalized for kw in keywords):
                return etype
        return None


# ──────────────────────────────────────────────────────────────
# MODULO 2A – RULE BASED DETECTOR
# ──────────────────────────────────────────────────────────────

class RuleBasedDetector:
    """Regex specializzate per dati italiani ed europei."""

    # (pattern, confidence)
    PATTERNS: Dict[EntityType, List[Tuple[str, float]]] = {
        EntityType.EMAIL: [
            # Quantificatori LIMITATi (RFC: local-part <=64, dominio <=255, TLD <=24).
            # Con '+' illimitato il pattern soffre di backtracking quadratico su
            # lunghe sequenze di caratteri ammessi senza '@' (es. "1.1.1.1..."),
            # sfruttabile come ReDoS/DoS su input non fidato. I limiti rendono il
            # lavoro per posizione costante senza perdere email reali.
            (r"\b[A-Za-z0-9._%+\-]{1,64}@[A-Za-z0-9.\-]{1,255}\.[A-Za-z]{2,24}\b", 0.99),
        ],
        EntityType.IBAN: [
            (r"\b[A-Z]{2}[0-9]{2}[A-Z0-9]{4}[0-9]{7}(?:[A-Z0-9]?){0,16}\b", 0.99),
        ],
        EntityType.CF: [
            # Codice Fiscale italiano (checksum verificato in detect()).
            (
                r"\b(?-i:[A-Z]{6}[0-9]{2}[ABCDEHLMPRST][0-9]{2}[A-Z][0-9]{3}[A-Z])\b",
                0.99,
            ),
        ],
        EntityType.PIVA: [
            # Preceduto da etichetta esplicita
            (r"(?:P\.?\s*IVA|Partita\s+IVA)[\s:]*(?P<v>[0-9]{11})", 0.99),
        ],
        EntityType.PHONE: [
            # Numeri italiani con prefisso internazionale
            (
                r"(?:\+39|0039)[\s\-\.]?(?:0[0-9]{1,3}|3[0-9]{2})[\s\-\.]?"
                r"[0-9]{3,4}[\s\-\.]?[0-9]{3,4}",
                0.92,
            ),
            # Numeri italiani senza prefisso
            (
                r"\b(?:0[0-9]{1,3}|3[0-9]{2})[\s\-\.]?[0-9]{3,4}[\s\-\.]?[0-9]{3,4}\b",
                0.85,
            ),
            # Numeri internazionali generici (non italiani): +1, +44, +33, ecc.
            (r"(?<!\d)\+(?!39\b|0039)[1-9][0-9]{6,14}(?!\d)", 0.80),
            # Internazionali con separatori (es. "+34 612 34 56 78", "+33 6 12 34 56 78",
            # "+49 30 12345678", "+351 912 345 678"): prefisso + 6-13 cifre con separatori.
            (r"(?<!\d)\+(?!39(?:[\s\-.]|\d))[1-9][0-9]{0,3}[\s\-.]?"
             r"[0-9](?:[\s\-.]?[0-9]){5,12}(?!\d)", 0.82),
        ],
        EntityType.CAP: [
            (r"\b[0-9]{5}\b", 0.65),
        ],
        EntityType.DATE: [
            # GG/MM/AAAA e varianti
            (
                r"\b(?:0?[1-9]|[12][0-9]|3[01])[/\-\.]"
                r"(?:0?[1-9]|1[0-2])[/\-\.](?:19|20)[0-9]{2}\b",
                0.95,
            ),
            # AAAA-MM-GG
            (
                r"\b(?:19|20)[0-9]{2}[/\-\.](?:0?[1-9]|1[0-2])[/\-\.]"
                r"(?:0?[1-9]|[12][0-9]|3[01])\b",
                0.95,
            ),
            # "15 gennaio 2023" / "15 gen 2023" / "15 gen. 2023" – italiano scritto
            # include "1°" ordinal marker e abbreviazioni con punto
            (
                r"\b(?:0?[1-9]|[12][0-9]|3[01])°?\s+"
                r"(?:gen(?:naio)?|feb(?:braio)?|mar(?:zo)?|apr(?:ile)?|mag(?:gio)?"
                r"|giu(?:gno)?|lug(?:lio)?|ago(?:sto)?|set(?:tembre)?|ott(?:obre)?"
                r"|nov(?:embre)?|dic(?:embre)?)\.?\s+(?:19|20)[0-9]{2}\b",
                0.90,
            ),
            # "giugno 2024" / "gen. 2024" – mese + anno senza giorno
            (
                r"\b(?:gen(?:naio)?|feb(?:braio)?|mar(?:zo)?|apr(?:ile)?|mag(?:gio)?"
                r"|giu(?:gno)?|lug(?:lio)?|ago(?:sto)?|set(?:tembre)?|ott(?:obre)?"
                r"|nov(?:embre)?|dic(?:embre)?)\.?\s+(?:19|20)[0-9]{2}\b",
                0.80,
            ),
        ],
        EntityType.AMOUNT: [
            (r"[€$£]\s*[0-9]{1,3}(?:\.[0-9]{3})*(?:,[0-9]{1,2})?", 0.95),
            (r"\b[0-9]{1,3}(?:\.[0-9]{3})*(?:,[0-9]{1,2})?\s*(?:euro|EUR|€)\b", 0.95),
            # Importi preceduti da etichetta esplicita senza simbolo valuta
            # (es. "Stipendio: 2.500,00" / "importo: 35000" / "reddito: 3.500,50")
            (
                r"(?:stipendio|salario|reddito|importo|totale|saldo|costo|prezzo|"
                r"fattura|quota|rata|acconto|rimborso|indennit[àa]|compenso|"
                r"salary|income|amount|balance|total|invoice|payment|"          # EN
                r"sueldo|ingreso|importe|precio|factura|"                        # ES
                r"salaire|revenu|montant|solde|facture|"                        # FR
                r"sal[áa]rio|rendimento|montante|pre[çc]o|fatura|"              # PT
                r"gehalt|betrag|summe|rechnung|einkommen)"                       # DE
                r"[\s:]*(?P<v>[0-9]{1,3}(?:\.[0-9]{3})*(?:,[0-9]{1,2})?)",
                0.82,
            ),
        ],
        EntityType.CARD: [
            # Luhn-validated credit card numbers (cifre contigue: Visa/MC/Amex/Diners/Discover)
            (r"\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13}"
             r"|3(?:0[0-5]|[68][0-9])[0-9]{11}|6(?:011|5[0-9]{2})[0-9]{12})\b", 0.95),
            # Numeri carta con separatori (spazio/trattino): "4111 1111 1111 1111".
            # Il Luhn (che ignora i non-cifra) filtra i falsi positivi.
            (r"(?<![0-9])(?:[0-9]{4}[ \-]){3}[0-9]{4}(?![0-9])", 0.90),
            (r"(?<![0-9])[0-9]{4}[ \-][0-9]{6}[ \-][0-9]{5}(?![0-9])", 0.90),  # Amex 15
        ],
        EntityType.BIC: [
            # BIC/SWIFT preceduto da etichetta
            (r"(?:BIC|SWIFT|BIC\s*/\s*SWIFT)[\s:]*"
             r"(?P<v>[A-Z]{6}[A-Z0-9]{2}(?:[A-Z0-9]{3})?)\b", 0.92),
            # BIC standalone (8 o 11 caratteri, 6 lettere + 2 alfanumerici + opz. 3).
            # (?-i:) forza il maiuscolo: senza, IGNORECASE matcherebbe parole minuscole.
            (r"\b(?-i:[A-Z]{6}[A-Z0-9]{2}(?:[A-Z0-9]{3})?)\b", 0.72),
        ],
        EntityType.CVV: [
            (r"(?:CVV2?|CVC2?|CID|security\s+code|codice\s+di\s+sicurezza|"
             r"c[óo]digo\s+de\s+seguridad|cryptogramme(?:\s+visuel)?|"
             r"c[óo]digo\s+de\s+seguran[çc]a|pr[üu]fziffer|sicherheitscode)"
             r"[\s:]*(?P<v>[0-9]{3,4})\b", 0.88),
        ],
        EntityType.CARD_EXPIRY: [
            (r"(?:scadenza|exp(?:iry|ires?)?(?:\s+date)?|valid\s+thru|good\s+thru|"
             r"vencimiento|caducidad|validit[ée]|expiration|validade|"
             r"g[üu]ltig\s+bis|ablaufdatum)"
             r"[\s:]*(?P<v>(?:0[1-9]|1[0-2])\s*/\s*(?:[0-9]{4}|[0-9]{2}))\b", 0.85),
        ],
        EntityType.ACCOUNT: [
            # Conto corrente italiano: preceduto da etichetta esplicita (ABI+CAB+cc = 12 cifre)
            (
                r"(?:c/c|conto\s+corrente|numero\s+conto|n[°\.]?\s*conto)[\s:]*"
                r"(?P<v>[0-9]{5}[\s\-]?[0-9]{5}[\s\-]?[0-9]{12}|[0-9]{12})",
                0.95,
            ),
            # Formato internazionale account preceduto da etichetta (IT/EN/ES/FR/PT/DE)
            (
                r"(?:account\s+(?:number|no\.?)|n[úu]mero\s+de\s+cuenta|"
                r"num[ée]ro\s+de\s+compte|n[úu]mero\s+de\s+conta|kontonummer)"
                r"[\s:]*(?P<v>[A-Z0-9]{6,20})",
                0.90,
            ),
        ],
        EntityType.ADDRESS: [
            # Indirizzi italiani: Via/Corso/Piazza/Viale + nome + numero civico
            # Spec target: 90% accuracy per indirizzi
            (
                r"\b(?:Via|V\.le|Viale|Corso|C\.so|Piazza|P\.(?:zza?)?|Vicolo|"
                r"Largo|Strada|Strada\s+Statale|S\.S\.|Loc(?:alità)?\.?|Borgata|"
                r"Contrada|Regione|Rione|Piazzale|Lungomare|Lungotevere)\s+"
                r"(?:dell[aeiou]?\s+|degli?\s+|dei?\s+|dello?\s+|d[i']?\s+)??"
                r"[A-ZÀ-Ùa-zà-ù][a-zà-ù]*(?:\s+[A-ZÀ-Ùa-zà-ù][a-zà-ù]*)?"
                r"(?:\s*,\s*|\s+)[0-9]+(?:\s*/\s*[A-Z0-9]+)?",
                0.78,
            ),
            # ES/FR/PT: prefisso stradale + nome + numero civico (numero dopo)
            (
                r"\b(?:Calle|C/|Avenida|Avda\.?|Plaza|Pza\.?|Paseo|Carrer|Camino|"
                r"Carretera|"                                                   # ES
                r"Rue|Avenue|Av\.?|Boulevard|Bd\.?|Place|Impasse|All[ée]e|"
                r"Chemin|Quai|Route|"                                           # FR
                r"Rua|Pra[çc]a|Travessa|Estrada|Alameda)\s+"                    # PT
                r"(?:d[aeiou]s?\s+|del\s+|de\s+la\s+|du\s+|des?\s+|do[s]?\s+)??"
                r"[A-ZÀ-ÖØ-Þa-zà-öø-ÿß][\wà-öø-ÿß]*(?:\s+[A-ZÀ-ÖØ-Þa-zà-öø-ÿß][\wà-öø-ÿß]*){0,3}"
                r"(?:\s*,\s*|\s+)(?:n[°ºo.]?\s*)?[0-9]+[A-Za-z]?(?:\s*/\s*[A-Z0-9]+)?",
                0.76,
            ),
            # DE: nome via con suffisso (-straße/-weg/-platz...) + numero civico
            (
                r"\b(?-i:[A-ZÄÖÜ][a-zäöüß]+(?:\s?[A-ZÄÖÜ][a-zäöüß]+)?"
                r"(?:stra(?:ße|sse)|str\.|weg|platz|gasse|allee|ring|damm|ufer))"
                r"\s+[0-9]+[a-z]?\b",
                0.76,
            ),
            # EN/US: numero civico + nome + tipo via (numero prima)
            (
                r"\b[0-9]{1,5}\s+[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){0,3}\s+"
                r"(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|"
                r"Dr|Court|Ct|Way|Square|Sq|Place|Pl|Terrace|Ter|Highway|Hwy)\b\.?",
                0.76,
            ),
        ],
        # Nomi di persone preceduti da titolo (IT/EN/ES/FR/PT/DE) o professionale.
        EntityType.PERSON: [
            (
                r"\b(?:Sig(?:\.ra?)?|Dott(?:\.ssa?)?|Dr|Prof(?:\.ssa?)?|"
                r"Avv|Ing|Arch|Rag|Geom|On|Sen|Gen|Col|"            # IT
                r"Mr|Mrs|Ms|Miss|Mx|Sir|Lord|Lady|"                  # EN
                r"Sr|Sra|Srta|Don|Do[ñn]a|"                          # ES/PT
                r"Mme|Mlle|"                                         # FR
                r"Herr|Frau)\.?\s+"                                  # DE
                r"(?-i:[A-ZÀ-ÖØ-Þ][a-zà-öø-ÿß]+(?:\s+[A-ZÀ-ÖØ-Þ][a-zà-öø-ÿß]+){0,3})",
                0.88,
            ),
            # FALLBACK basato sul contesto: nome proprio (Nome Cognome) preceduto
            # da un'etichetta tipica in 6 lingue. Rete di sicurezza quando il NER
            # non è disponibile (es. "Cliente: Mario Rossi", "Nom : Jean Dupont").
            (
                r"(?:"
                # IT
                r"cliente|dipendente|paziente|intestatario|titolare|assicurato|"
                r"beneficiario|contatto(?:\s+(?:emergenza|di\s+emergenza))?|"
                r"referente|destinatario|mittente|richiedente|utente|nome|"
                # EN
                r"name|full\s+name|customer|employee|patient|client|contact|"
                r"holder|policyholder|beneficiary|recipient|sender|applicant|"
                # ES
                r"nombre|empleado|titular|contacto|asegurado|usuario|remitente|"
                # FR
                r"nom(?:\s+complet)?|employ[ée]|titulaire|b[ée]n[ée]ficiaire|"
                r"destinataire|exp[ée]diteur|utilisateur|assur[ée]|"
                # PT
                r"funcion[áa]rio|contato|benefici[áa]rio|destinat[áa]rio|"
                r"remetente|utilizador|segurado|"
                # DE
                r"kunde|mitarbeiter|inhaber|kontakt|empf[äa]nger|absender|"
                r"benutzer|versicherter|name"
                r")"
                r"[ \t]*:?[ \t]*"
                r"(?P<v>(?-i:[A-ZÀ-ÖØ-Þ][a-zà-öø-ÿß'’]+(?:[ \t]+(?:de[il]?[ \t]+|"
                r"d[a'][ \t]*|van[ \t]+|von[ \t]+|della?[ \t]+|do[s]?[ \t]+)?"
                r"[A-ZÀ-ÖØ-Þ][a-zà-öø-ÿß'’]+){1,3}))",
                0.75,
            ),
        ],
        # ── Documenti d'identità (IT/EN/ES/FR/PT/DE) ──
        EntityType.PASSPORT: [
            (r"(?:passaporto|passport|pasaporte|passeport|passaporte|"
             r"reisepass|pass)(?:\s*(?:n[°ºo.]?|no\.?|nr\.?|num(?:ero|éro)?|number|#))?"
             r"[\s:]*(?P<v>[A-Z]{1,2}\s?[0-9]{6,8})\b", 0.90),
        ],
        EntityType.ID_CARD: [
            (r"(?:carta\s+d['’i]?\s*identit[àa]|carta\s+identit[àa]|"
             r"documento\s+d['’i]?\s*identit[àa]|identity\s+card|id\s+card|"
             r"national\s+id|carte\s+nationale\s+d['’]identit[ée]|cni|"
             r"cart[ãa]o\s+de\s+cidad[ãa]o|bilhete\s+de\s+identidade|"
             r"personalausweis|ausweis(?:nummer)?|"
             r"documento\s+nacional|dni|nie)"
             r"(?:\s*(?:n[°ºo.]?|no\.?|nr\.?|num(?:ero|éro)?|number|#))?"
             r"[\s:]*(?P<v>[A-Z]{0,2}\s?[0-9]{6,9}[ \-]?[A-Z]?)\b", 0.88),
            # DNI spagnolo standalone (8 cifre + lettera di controllo)
            (r"\b(?-i:[0-9]{8}[ \-]?[A-Z])\b", 0.72),
            # NIE spagnolo standalone (X/Y/Z + 7 cifre + lettera)
            (r"\b(?-i:[XYZ][ \-]?[0-9]{7}[ \-]?[A-Z])\b", 0.80),
        ],
        EntityType.DRIVER_LICENSE: [
            (r"(?:patente(?:\s+di\s+guida)?|driver'?s?\s+licen[cs]e|"
             r"driving\s+licen[cs]e|permiso\s+de\s+conducir|permis\s+de\s+conduire|"
             r"carta\s+de\s+condu[çc][ãa]o|f[üu]hrerschein)"
             r"(?:\s*(?:n[°ºo.]?|no\.?|nr\.?|num(?:ero|éro)?|number|#))?"
             r"[\s:]*(?P<v>[A-Z]{1,2}[0-9]{6,10}[A-Z]?)\b", 0.85),
        ],
        EntityType.LICENSE_PLATE: [
            (r"(?:targa|license\s+plate|number\s+plate|matr[íi]cula|"
             r"plaque(?:\s+d['’]immatriculation)?|matr[íi]cula|"
             r"matr[íi]cula|kennzeichen)[\s:]*"
             r"(?P<v>(?-i:[A-Z]{1,3}[ \-]?[0-9]{1,4}[ \-]?[A-Z]{0,3}[0-9]{0,4}))\b", 0.82),
        ],
        # ── Identificativi fiscali (US/UK/ES/FR/PT/DE + EU VAT) ──
        EntityType.TAX_ID: [
            # US SSN
            (r"(?:SSN|social\s+security(?:\s+(?:no|number))?)[\s.:#]*"
             r"(?P<v>[0-9]{3}-?[0-9]{2}-?[0-9]{4})\b", 0.90),
            (r"\b[0-9]{3}-[0-9]{2}-[0-9]{4}\b", 0.78),
            # UK National Insurance Number (case-sensitive: evita match su minuscole)
            (r"\b(?-i:[A-CEGHJ-PR-TW-Z]{2}\s?[0-9]{2}\s?[0-9]{2}\s?[0-9]{2}\s?[A-D])\b", 0.82),
            # EU VAT preceduto da etichetta (formato generico paese+cifre)
            (r"(?:VAT(?:\s+(?:no|number|reg(?:istration)?))?|tax\s+id|"
             r"steuer-?id|steueridentifikationsnummer|steuernummer|"
             r"num[ée]ro\s+fiscal|n[íi]f|contribuinte)[\s.:#]*"
             r"(?P<v>[A-Z]{0,2}[0-9A-Z]{8,13})\b", 0.85),
            # FR numéro de sécurité sociale / NIR (13 cifre + 2 di controllo)
            (r"(?:s[ée]curit[ée]\s+sociale|num[ée]ro\s+de\s+s[ée]curit[ée]|nir|"
             r"n[°º]?\s*s[ée]cu)[\s.:#]*"
             r"(?P<v>[12][ ]?[0-9]{2}[ ]?[0-9]{2}[ ]?[0-9AB]{2}[ ]?[0-9]{3}"
             r"[ ]?[0-9]{3}(?:[ ]?[0-9]{2})?)\b", 0.88),
            # ES número de la seguridad social (NUSS) / FR generico etichettato
            (r"(?:seguridad\s+social|n[úu]mero\s+de\s+afiliaci[óo]n|nuss|"
             r"sozialversicherungsnummer|seguran[çc]a\s+social|niss)[\s.:#]*"
             r"(?P<v>[0-9]{2}[ /\-]?[0-9]{6,12})\b", 0.85),
            # PT NIF / ES NIF-CIF standalone con etichetta breve già coperti sopra
        ],
        # ── Identificativi di rete ──
        EntityType.IP_ADDRESS: [
            # IPv4
            (r"\b(?:(?:25[0-5]|2[0-4][0-9]|1?[0-9]?[0-9])\.){3}"
             r"(?:25[0-5]|2[0-4][0-9]|1?[0-9]?[0-9])\b", 0.85),
            # IPv6 (forma piena o compressa con '::')
            (r"\b(?:[A-F0-9]{1,4}:){7}[A-F0-9]{1,4}\b", 0.85),
            (r"(?<![:.\w])(?:[A-F0-9]{1,4}:){1,6}:(?:[A-F0-9]{1,4}:?){0,5}[A-F0-9]{1,4}", 0.80),
        ],
        EntityType.MAC_ADDRESS: [
            # confidence > IPv6 così, in caso di span identico, vince MAC
            (r"\b(?:[0-9A-F]{2}[:\-]){5}[0-9A-F]{2}\b", 0.92),
        ],
        EntityType.URL: [
            (r"\bhttps?://[^\s<>\"')\]]+", 0.85),
            (r"\bwww\.[A-Za-z0-9\-]+\.[A-Za-z]{2,24}(?:/[^\s<>\"')\]]*)?", 0.80),
        ],
        EntityType.USERNAME: [
            # @handle (non email): preceduto da inizio/spazio, non da carattere parola
            (r"(?<![\w@.])@[A-Za-z0-9_]{2,30}\b", 0.62),
            (r"(?:username|user(?:name)?|handle|nickname|account|login|utente|"
             r"usuario|utilisateur|identifiant|utilizador|benutzer(?:name)?|nutzername)"
             r"[\s:]*(?P<v>[A-Za-z0-9._\-]{3,30})\b", 0.70),
        ],
        # ── Identificativi lavorativi e sanitari (IT/EN/ES/FR/PT/DE) ──
        EntityType.EMPLOYEE_ID: [
            (r"(?:matricola(?:\s+aziendale)?|employee\s+(?:id|number|no\.?)|"
             r"badge\s+(?:id|number|no\.?)|emp(?:loyee)?\s*id|staff\s+(?:id|number)|"
             r"n[úu]mero\s+de\s+empleado|matricule|num[ée]ro\s+d['’]employ[ée]|"
             r"n[úu]mero\s+de\s+funcion[áa]rio|personalnummer|mitarbeiternummer)"
             r"[\s.:#]*(?P<v>[A-Z]{0,5}[-/]?[0-9]{3,10})\b", 0.85),
        ],
        EntityType.MEDICAL_ID: [
            (r"(?:cartella\s+clinica|medical\s+record(?:\s+(?:no|number))?|"
             r"patient\s+(?:id|number|no\.?)|id\s+paziente|"
             r"codice\s+(?:paziente|sanitario)|health\s+id|"
             r"n[úu]mero\s+de\s+(?:paciente|historia\s+cl[íi]nica)|"
             r"dossier\s+m[ée]dical|num[ée]ro\s+de\s+patient|"
             r"n[úu]mero\s+de\s+(?:utente|processo\s+cl[íi]nico)|"
             r"patientennummer|krankenakte)"
             r"[\s.:#]*(?P<v>[A-Z]{0,5}[-/]?[0-9]{4,12})\b", 0.85),
            (r"\bMED[-/]?[0-9]{4,12}\b", 0.78),
        ],
        EntityType.INSURANCE_ID: [
            (r"(?:assicurazione(?:\s+sanitaria)?|polizza|health\s+insurance|"
             r"insurance(?:\s+(?:no|number|id|policy))?|policy\s+(?:no|number)|"
             r"p[óo]liza|seguro|num[ée]ro\s+(?:de\s+)?"
             r"(?:police|d['’]assurance)|police\s+d['’]assurance|ap[óo]lice|"
             r"versicherungsnummer|versicherung)"
             r"[\s.:#nº°]*(?P<v>[A-Z]{0,5}[-/]?[0-9]{6,14})\b", 0.85),
            (r"\bINS[-/]?[0-9]{6,14}\b", 0.78),
        ],
    }

    # Codice Fiscale: valori per caratteri in posizione dispari (1-indexed)
    _CF_ODD: Dict[str, int] = {
        **{str(i): v for i, v in enumerate([1,0,5,7,9,13,15,17,19,21])},
        **dict(zip("ABCDEFGHIJKLMNOPQRSTUVWXYZ",
                   [1,0,5,7,9,13,15,17,19,21,2,4,18,20,11,3,6,8,12,14,16,10,22,25,24,23]))
    }
    # Codice Fiscale: valori per caratteri in posizione pari (1-indexed) = ordinale 0-25 per lettere
    _CF_EVEN: Dict[str, int] = {
        **{str(i): i for i in range(10)},
        **{c: i for i, c in enumerate("ABCDEFGHIJKLMNOPQRSTUVWXYZ")}
    }

    def detect(
        self,
        text: str,
        threshold: float = 0.7,
        entity_types: Optional[List[EntityType]] = None,
    ) -> List[DetectedEntity]:
        results: List[DetectedEntity] = []
        allowed = set(entity_types) if entity_types else None

        for entity_type, patterns in self.PATTERNS.items():
            if allowed is not None and entity_type not in allowed:
                continue
            for pattern, confidence in patterns:
                if confidence < threshold:
                    continue
                for m in re.finditer(pattern, text, re.IGNORECASE):
                    # Se il pattern definisce un gruppo nominato "v", si tokenizza
                    # SOLO quel sotto-span (il valore sensibile) lasciando intatta
                    # l'etichetta di contesto (es. "Passaporto: " resta in chiaro,
                    # viene sostituito solo "YA7654321"). Altrimenti si usa l'intero
                    # match. Questo evita di cancellare le etichette dal documento.
                    if "v" in m.groupdict() and m.group("v") is not None:
                        value = m.group("v")
                        span_start, span_end = m.span("v")
                    else:
                        value = m.group(0)
                        span_start, span_end = m.start(), m.end()
                    if entity_type == EntityType.CARD and not self._luhn(value):
                        continue
                    if entity_type == EntityType.IBAN and not self._validate_iban(value):
                        continue
                    if entity_type == EntityType.CF and not self._validate_cf(value):
                        continue
                    if entity_type == EntityType.PIVA and not self._validate_piva(value):
                        continue
                    results.append(DetectedEntity(
                        value=value,
                        type=entity_type,
                        confidence=confidence,
                        start=span_start,
                        end=span_end,
                        source="rule",
                    ))

        return results

    @staticmethod
    def _luhn(number: str) -> bool:
        digits = [int(c) for c in number if c.isdigit()]
        if len(digits) < 13:
            return False
        total = 0
        for i, d in enumerate(reversed(digits)):
            if i % 2 == 1:
                d *= 2
                if d > 9:
                    d -= 9
            total += d
        return total % 10 == 0

    @staticmethod
    def _validate_iban(iban: str) -> bool:
        """Verifica IBAN con algoritmo modulo 97 (ISO 13616)."""
        clean = re.sub(r"\s+", "", iban).upper()
        if len(clean) < 5:
            return False
        rearranged = clean[4:] + clean[:4]
        numeric = "".join(str(ord(c) - 55) if c.isalpha() else c for c in rearranged)
        try:
            return int(numeric) % 97 == 1
        except ValueError:
            return False

    @staticmethod
    def _validate_piva(text: str) -> bool:
        """Verifica il checksum della Partita IVA italiana (11 cifre)."""
        digits = [int(c) for c in text if c.isdigit()]
        if len(digits) != 11:
            return False
        total = 0
        for i, d in enumerate(digits[:10]):
            if i % 2 == 0:
                total += d
            else:
                doubled = d * 2
                total += doubled if doubled < 10 else doubled - 9
        expected = (10 - (total % 10)) % 10
        return digits[10] == expected

    def _validate_cf(self, cf: str) -> bool:
        """Verifica il carattere di controllo del Codice Fiscale italiano."""
        cf = cf.upper()
        if len(cf) != 16:
            return False
        total = sum(
            self._CF_ODD.get(c, 0) if i % 2 == 0 else self._CF_EVEN.get(c, 0)
            for i, c in enumerate(cf[:15])
        )
        return chr(total % 26 + ord("A")) == cf[15]


# ──────────────────────────────────────────────────────────────
# MODULO 2B – NER DETECTOR
# ──────────────────────────────────────────────────────────────

class NERDetector:
    """NER con spaCy, GLiNER e/o Microsoft Presidio (tutti opzionali, graceful degradation)."""

    _SPACY_LABEL_MAP = {
        "PER": EntityType.PERSON,
        "PERSON": EntityType.PERSON,
        "ORG": EntityType.ORGANIZATION,
        "LOC": EntityType.LOCATION,
        "GPE": EntityType.LOCATION,
        "FAC": EntityType.ADDRESS,
    }

    _PRESIDIO_LABEL_MAP = {
        "PERSON": EntityType.PERSON,
        "PHONE_NUMBER": EntityType.PHONE,
        "EMAIL_ADDRESS": EntityType.EMAIL,
        "IBAN_CODE": EntityType.IBAN,
        "CREDIT_CARD": EntityType.CARD,
        "LOCATION": EntityType.LOCATION,
        "ORGANIZATION": EntityType.ORGANIZATION,
        "DATE_TIME": EntityType.DATE,
        "IT_FISCAL_CODE": EntityType.CF,
        # Entità Presidio aggiuntive (US/UK/generiche)
        "IP_ADDRESS": EntityType.IP_ADDRESS,
        "URL": EntityType.URL,
        "US_SSN": EntityType.TAX_ID,
        "US_ITIN": EntityType.TAX_ID,
        "US_PASSPORT": EntityType.PASSPORT,
        "US_DRIVER_LICENSE": EntityType.DRIVER_LICENSE,
        "US_BANK_NUMBER": EntityType.ACCOUNT,
        "UK_NHS": EntityType.MEDICAL_ID,
        "UK_NINO": EntityType.TAX_ID,
        "MEDICAL_LICENSE": EntityType.MEDICAL_ID,
        "IT_DRIVER_LICENSE": EntityType.DRIVER_LICENSE,
        "IT_PASSPORT": EntityType.PASSPORT,
        "IT_IDENTITY_CARD": EntityType.ID_CARD,
        "IT_VAT_CODE": EntityType.PIVA,
        "CRYPTO": EntityType.ACCOUNT,
    }

    # GLiNER usa etichette in linguaggio naturale (zero-shot)
    _GLINER_LABELS = [
        "person", "organization", "location", "address",
        "email", "phone number", "IBAN", "fiscal code",
        "date of birth", "monetary amount",
        "credit card number", "bank account number",
        "passport number", "identity card number", "driver license number",
        "license plate", "social security number", "VAT number",
        "IP address", "MAC address", "URL", "username",
        "employee id", "medical record number", "insurance policy number",
    ]
    _GLINER_LABEL_MAP = {
        "person": EntityType.PERSON,
        "organization": EntityType.ORGANIZATION,
        "location": EntityType.LOCATION,
        "address": EntityType.ADDRESS,
        "email": EntityType.EMAIL,
        "phone number": EntityType.PHONE,
        "iban": EntityType.IBAN,
        "fiscal code": EntityType.CF,
        "date of birth": EntityType.DATE,
        "monetary amount": EntityType.AMOUNT,
        "credit card number": EntityType.CARD,
        "bank account number": EntityType.ACCOUNT,
        "passport number": EntityType.PASSPORT,
        "identity card number": EntityType.ID_CARD,
        "driver license number": EntityType.DRIVER_LICENSE,
        "license plate": EntityType.LICENSE_PLATE,
        "social security number": EntityType.TAX_ID,
        "vat number": EntityType.TAX_ID,
        "ip address": EntityType.IP_ADDRESS,
        "mac address": EntityType.MAC_ADDRESS,
        "url": EntityType.URL,
        "username": EntityType.USERNAME,
        "employee id": EntityType.EMPLOYEE_ID,
        "medical record number": EntityType.MEDICAL_ID,
        "insurance policy number": EntityType.INSURANCE_ID,
    }

    # Parole-etichetta che i modelli NER a volte scambiano per PERSON/LOC/ORG
    # (IT + EN). Non sono mai dati personali: vengono scartate in _sanitize.
    _NER_STOPWORDS = frozenset({
        "attenzione", "cliente", "dipendente", "paziente", "contatto",
        "relazione", "scadenza", "note", "documento", "indirizzo", "telefono",
        "matricola", "assicurazione", "polizza", "passaporto", "patente",
        "targa", "email", "data", "nascita", "intestatario", "titolare",
        "beneficiario", "referente", "destinatario", "mittente", "richiedente",
        "name", "customer", "employee", "patient", "contact", "relationship",
        "expiry", "note", "notes", "address", "phone", "passport", "license",
        "insurance", "policy", "holder", "beneficiary", "recipient", "sender",
        # ES
        "atención", "atencion", "nombre", "apellido", "empleado", "asegurado",
        "titular", "contacto", "dirección", "direccion", "teléfono", "telefono",
        "fecha", "nacimiento", "usuario", "póliza", "poliza", "seguro", "cuenta",
        # FR
        "attention", "nom", "prénom", "prenom", "client", "employé", "employe",
        "assuré", "assure", "adresse", "téléphone", "telephone", "naissance",
        "contact", "destinataire", "expéditeur", "expediteur", "utilisateur",
        # PT
        "atenção", "atencao", "cliente", "funcionário", "funcionario", "morada",
        "telefone", "nascimento", "segurado", "apólice", "apolice", "utilizador",
        # DE
        "achtung", "kunde", "mitarbeiter", "patient", "anschrift", "telefon",
        "geburtsdatum", "versicherter", "versicherung", "benutzer", "konto",
        "name", "vorname", "nachname",
    })

    # Articoli/determinativi (IT/EN/ES/FR/PT/DE) da rifilare in testa/coda agli
    # span NER di PERSON/ORG/LOCATION (i modelli statistici li agganciano spesso).
    _ARTICLES = frozenset({
        "il", "lo", "la", "i", "gli", "le", "l", "un", "uno", "una",       # IT
        "the", "a", "an",                                                   # EN
        "el", "los", "las", "unos", "unas",                                 # ES
        "les", "des", "du", "de",                                           # FR/IT/ES/PT
        "o", "os", "as", "um", "uma",                                       # PT
        "der", "die", "das", "den", "dem", "ein", "eine", "einen",         # DE
    })

    # Tipi NER su cui applicare il trim di articoli iniziali e token minuscoli
    # (verbi/preposizioni) finali: il "core" del nome è fatto di token Maiuscoli.
    _TRIMMABLE_TYPES = frozenset({
        EntityType.PERSON, EntityType.ORGANIZATION, EntityType.LOCATION,
    })

    # Un gruppo di modelli per lingua (in ordine di preferenza lg→md→sm). Viene
    # caricato il primo disponibile di ciascun gruppo: così la detection è
    # multilingue (IT, EN, ES, FR, PT, DE) e degrada con grazia se un modello
    # non è installato. xx_ent_wiki_sm è un fallback multilingue.
    _SPACY_MODEL_GROUPS = (
        ("it_core_news_lg", "it_core_news_md", "it_core_news_sm"),
        ("en_core_web_lg", "en_core_web_md", "en_core_web_sm"),
        ("es_core_news_lg", "es_core_news_md", "es_core_news_sm"),
        ("fr_core_news_lg", "fr_core_news_md", "fr_core_news_sm"),
        ("pt_core_news_lg", "pt_core_news_md", "pt_core_news_sm"),
        ("de_core_news_lg", "de_core_news_md", "de_core_news_sm"),
        ("xx_ent_wiki_sm",),
    )

    def __init__(self) -> None:
        self._spacy_nlp = None        # modello primario (retrocompatibilità)
        self._spacy_extra = []        # modelli aggiuntivi (altre lingue)
        self._spacy_models = []       # nomi modelli caricati (per config Presidio)
        self._presidio = None
        self._presidio_langs = ["en"]
        self._gliner = None
        # Lazy-loading: i modelli (6 spaCy + Presidio + GLiNER) sono pesanti da
        # caricare (~30-40s). Si caricano alla prima detect(), così costruire un
        # NERDetector è istantaneo e, se il NER non viene mai usato, non si paga
        # nulla. _loaded=False ⇒ "da caricare"; assente (creazione via __new__,
        # es. nei test) ⇒ "configurato a mano, non caricare".
        self._loaded = False

    def _ensure_loaded(self) -> None:
        if getattr(self, "_loaded", None) is not False:
            return  # già caricato (True) oppure istanza __new__ configurata a mano
        self._loaded = True
        self._load_spacy()
        self._load_presidio()
        self._load_gliner()

    def _load_spacy(self) -> None:
        try:
            import spacy
        except ImportError:
            logger.warning("spaCy non disponibile: pip install spacy")
            return
        loaded: List[str] = []
        for group in self._SPACY_MODEL_GROUPS:
            for model in group:
                try:
                    nlp = spacy.load(model)
                except (OSError, IOError):
                    continue
                if self._spacy_nlp is None:
                    self._spacy_nlp = nlp
                else:
                    self._spacy_extra.append(nlp)
                loaded.append(model)
                break  # primo disponibile del gruppo (lingua)
        self._spacy_models = loaded
        if loaded:
            logger.info(f"spaCy: modelli caricati: {', '.join(loaded)}")
        else:
            logger.warning(
                "Nessun modello spaCy trovato. Installare ad es.:\n"
                "  python -m spacy download it_core_news_sm\n"
                "  python -m spacy download en_core_web_sm  (es/fr/pt/de analoghi)"
            )

    def _load_presidio(self) -> None:
        try:
            from presidio_analyzer import AnalyzerEngine
        except ImportError:
            logger.warning("Presidio non disponibile: pip install presidio-analyzer")
            return
        # Costruisce un motore NLP multilingue riusando i modelli spaCy installati
        # (IT/EN/ES/FR/PT/DE). Senza questa config, AnalyzerEngine supporta solo
        # l'inglese. Su qualsiasi errore si ricade sul default (en).
        models_cfg, langs = [], []
        for name in self._spacy_models:
            lang = name[:2]
            if lang == "xx" or lang in langs:
                continue
            langs.append(lang)
            models_cfg.append({"lang_code": lang, "model_name": name})
        if models_cfg:
            try:
                from presidio_analyzer.nlp_engine import NlpEngineProvider
                engine = NlpEngineProvider(nlp_configuration={
                    "nlp_engine_name": "spacy", "models": models_cfg,
                }).create_engine()
                self._presidio = AnalyzerEngine(
                    nlp_engine=engine, supported_languages=langs)
                self._presidio_langs = langs
                logger.info(f"Presidio multilingue: {', '.join(langs)}")
                return
            except Exception as exc:
                logger.warning(f"Presidio multilingue non configurato ({exc}); default 'en'")
        try:
            self._presidio = AnalyzerEngine()
            self._presidio_langs = ["en"]
            logger.info("Presidio analyzer inizializzato (en)")
        except Exception as exc:
            logger.warning(f"Presidio non inizializzato: {exc}")

    def _load_gliner(self) -> None:
        try:
            from gliner import GLiNER
            # Modello multilingue leggero, buono per italiano
            self._gliner = GLiNER.from_pretrained("urchade/gliner_multi-v2.1")
            logger.info("GLiNER: modello 'urchade/gliner_multi-v2.1' caricato")
        except ImportError:
            logger.warning("GLiNER non disponibile: pip install gliner")
        except Exception as exc:
            logger.warning(f"GLiNER non caricato: {exc}")

    def detect(
        self,
        text: str,
        threshold: float = 0.7,
        entity_types: Optional[List[EntityType]] = None,
    ) -> List[DetectedEntity]:
        self._ensure_loaded()
        entities: List[DetectedEntity] = []

        if self._spacy_nlp:
            entities.extend(self._spacy_detect(text))

        if self._gliner:
            entities.extend(self._gliner_detect(text, threshold))

        if self._presidio:
            entities.extend(self._presidio_detect(text, threshold))

        # Sanificazione span NER: i modelli statistici a volte estendono lo span
        # oltre l'a-capo o etichettano le intestazioni di campo (es. "ATTENZIONE:",
        # "CVV:"). Si tagliano gli a-capo e si scartano le label strutturali.
        entities = [s for e in entities if (s := self._sanitize(text, e)) is not None]

        # Filtro soglia uniforme: spaCy assegna una confidence fissa (0.85) e non
        # filtra internamente come fanno GLiNER e Presidio. Applicarlo qui rende il
        # comportamento di --confidence coerente per tutte le sorgenti NER (con
        # soglia > 0.85 le entità spaCy vengono correttamente scartate).
        entities = [e for e in entities if e.confidence >= threshold]

        if entity_types is not None:
            allowed = set(entity_types)
            entities = [e for e in entities if e.type in allowed]

        return entities

    @staticmethod
    def _sanitize(text: str, ent: "DetectedEntity") -> Optional["DetectedEntity"]:
        """Ripulisce uno span NER. Ritorna None se va scartato.

        - Tronca al primo a-capo (un dato PII non attraversa righe diverse).
        - Rimuove spazi iniziali/finali ricalcolando gli offset.
        - Scarta le intestazioni di campo: valore seguito (saltati gli spazi) da
          ':' (es. "ATTENZIONE:", "IBAN:", "CVV:") — sono label, non valori.
        - Scarta i token singoli tutto-maiuscolo (acronimi come "BIC"): i nomi
          propri PII sono normalmente a maiuscola/minuscola mista.
        """
        s, e = ent.start, ent.end
        segment = text[s:e]
        nl = segment.find("\n")
        if nl != -1:
            e = s + nl
            segment = text[s:e]
        lead = len(segment) - len(segment.lstrip())
        trail = len(segment) - len(segment.rstrip())
        s, e = s + lead, e - trail
        value = text[s:e]
        if not value:
            return None
        # Label con ':' subito dopo (saltati gli spazi) o ':' incluso nello span.
        j = e
        while j < len(text) and text[j] in " \t":
            j += 1
        if (j < len(text) and text[j] == ":") or value.endswith(":"):
            return None
        # Acronimi tutto-maiuscolo (es. "BIC") e parole-etichetta comuni: rumore NER.
        if " " not in value and value.isupper() and value.isalpha():
            return None
        if value.lower() in NERDetector._NER_STOPWORDS:
            return None
        # Trim del rumore spaCy su PERSON/ORG/LOCATION: rimuove in testa gli
        # articoli/etichette e in coda i token tutto-minuscolo (verbi/preposizioni
        # agganciati), mantenendo il nucleo Maiuscolo. Es.: "El paciente Antonio
        # Fernández acudió" → "Antonio Fernández", "la consulta en Sevilla" → "Sevilla".
        if ent.type in NERDetector._TRIMMABLE_TYPES and " " in value:
            toks = list(re.finditer(r"\S+", value))
            lo, hi = 0, len(toks) - 1
            def _drop(tok: str) -> bool:
                w = tok.strip(".,;:'’").lower()
                return tok.islower() or w in NERDetector._ARTICLES or w in NERDetector._NER_STOPWORDS
            while lo <= hi and _drop(toks[lo].group()):
                lo += 1
            while hi >= lo and _drop(toks[hi].group()):
                hi -= 1
            if lo > hi:
                return None
            new_s = s + toks[lo].start()
            new_e = s + toks[hi].end()
            s, e = new_s, new_e
            value = text[s:e]
            if not value:
                return None
        return DetectedEntity(
            value=value,
            type=ent.type,
            confidence=ent.confidence,
            start=s,
            end=e,
            source=ent.source,
        )

    def _spacy_detect(self, text: str) -> List[DetectedEntity]:
        results = []
        # Modello primario + eventuali modelli di altre lingue. getattr difende
        # i test che istanziano NERDetector via __new__ senza _spacy_extra.
        models = [self._spacy_nlp] + list(getattr(self, "_spacy_extra", []))
        for nlp in models:
            if nlp is None:
                continue
            for ent in nlp(text).ents:
                etype = self._SPACY_LABEL_MAP.get(ent.label_)
                if etype:
                    results.append(DetectedEntity(
                        value=ent.text,
                        type=etype,
                        confidence=0.85,
                        start=ent.start_char,
                        end=ent.end_char,
                        source="spacy",
                    ))
        return results

    def _gliner_detect(self, text: str, threshold: float) -> List[DetectedEntity]:
        results = []
        try:
            # GLiNER lavora meglio su testi <= 512 token; spezza se necessario
            hits = self._gliner.predict_entities(
                text, self._GLINER_LABELS, threshold=threshold
            )
            for h in hits:
                label_key = h["label"].lower()
                etype = self._GLINER_LABEL_MAP.get(label_key)
                if etype:
                    results.append(DetectedEntity(
                        value=h["text"],
                        type=etype,
                        confidence=float(h.get("score", 0.8)),
                        start=h["start"],
                        end=h["end"],
                        source="gliner",
                    ))
        except Exception as exc:
            logger.warning(f"GLiNER error: {exc}")
        return results

    def _presidio_detect(self, text: str, threshold: float) -> List[DetectedEntity]:
        results = []
        entities_to_check = list(self._PRESIDIO_LABEL_MAP.keys())
        seen: set = set()
        for lang in getattr(self, "_presidio_langs", ["en"]):
            try:
                hits = self._presidio.analyze(
                    text=text, language=lang, entities=entities_to_check
                )
                for h in hits:
                    if h.score < threshold:
                        continue
                    key = (h.start, h.end, h.entity_type)
                    if key in seen:
                        continue
                    seen.add(key)
                    etype = self._PRESIDIO_LABEL_MAP.get(h.entity_type)
                    if etype:
                        results.append(DetectedEntity(
                            value=text[h.start:h.end],
                            type=etype,
                            confidence=h.score,
                            start=h.start,
                            end=h.end,
                            source="presidio",
                        ))
            except Exception as exc:
                logger.debug(f"Presidio [{lang}]: {exc}")
        return results


# ──────────────────────────────────────────────────────────────
# MODULO 2 – DETECTION ENGINE (ibrido)
# ──────────────────────────────────────────────────────────────

class SensitiveDataDetectionEngine:
    """Motore ibrido: regex + NER con confidence score configurabile."""

    def __init__(
        self,
        confidence_threshold: float = 0.7,
        use_ner: bool = True,
        entity_types: Optional[List[EntityType]] = None,
    ) -> None:
        self.threshold = confidence_threshold
        self.entity_types = entity_types  # None = tutti i tipi
        self._rule = RuleBasedDetector()
        self._ner = NERDetector() if use_ner else None

    def detect(self, text: str) -> List[DetectedEntity]:
        entities = self._rule.detect(text, self.threshold, self.entity_types)
        if self._ner:
            entities.extend(self._ner.detect(text, self.threshold, self.entity_types))
        return self._deduplicate(entities)

    @staticmethod
    def _boost_confidence(entities: List[DetectedEntity]) -> List[DetectedEntity]:
        """Aumenta confidence +0.05 (max 0.99) quando più sorgenti concordano sullo stesso span."""
        # Raggruppa per (start, end, type): se le sorgenti sono diverse, boost
        from collections import defaultdict
        span_sources: Dict[Tuple[int, int, EntityType], set] = defaultdict(set)
        for e in entities:
            span_sources[(e.start, e.end, e.type)].add(e.source)

        boosted: List[DetectedEntity] = []
        for e in entities:
            key = (e.start, e.end, e.type)
            if len(span_sources[key]) > 1:
                boosted.append(DetectedEntity(
                    value=e.value,
                    type=e.type,
                    confidence=min(0.99, e.confidence + 0.05),
                    start=e.start,
                    end=e.end,
                    source=e.source,
                ))
            else:
                boosted.append(e)
        return boosted

    @staticmethod
    def _deduplicate(entities: List[DetectedEntity]) -> List[DetectedEntity]:
        """Boost confidence multi-sorgente, rimuove overlap, ordina per posizione.

        In caso di sovrapposizione si mantiene lo span che copre PIÙ caratteri;
        a parità di lunghezza vince la confidence maggiore. Questa scelta è
        deliberata per un tool di privacy: sostituire uno span lungo con uno più
        corto annidato (anche se a confidence superiore) lascerebbe esposti i
        caratteri non più coperti — una potenziale fuga di dati sensibili.
        Esempio: ADDRESS "Via Roma 1, 00100" (0-17) che contiene CAP "00100"
        (12-17): mantenendo il CAP, "Via Roma 1, " resterebbe in chiaro.
        """
        entities = SensitiveDataDetectionEngine._boost_confidence(entities)
        # start crescente → span più lungo → confidence più alta.
        entities.sort(key=lambda e: (e.start, -(e.end - e.start), -e.confidence))
        result: List[DetectedEntity] = []
        last_end = -1
        for ent in entities:
            if ent.start >= last_end:
                result.append(ent)
                last_end = ent.end
            # Overlap: l'entità inizia dentro l'ultimo span già tenuto. La si
            # scarta perché l'ordinamento per lunghezza decrescente garantisce
            # che lo span trattenuto copra almeno fin dove copriva questa; tenerla
            # ridurrebbe la copertura o introdurrebbe span sovrapposti (che
            # corromperebbero la tokenizzazione basata sugli indici).
        return result


# ──────────────────────────────────────────────────────────────
# MODULO 3 – TOKENIZATION ENGINE
# ──────────────────────────────────────────────────────────────

class TokenizationEngine:
    """
    Sostituisce i dati sensibili con placeholder deterministici.
    Regola: stesso valore → stesso token (sempre).
    """

    def __init__(self) -> None:
        self._value_to_token: Dict[str, str] = {}
        self._counters: Dict[str, int] = {}

    def reset(self) -> None:
        self._value_to_token.clear()
        self._counters.clear()

    def tokenize(
        self, text: str, entities: List[DetectedEntity]
    ) -> Tuple[str, Dict[str, str]]:
        """
        Restituisce (testo_anonimizzato, mappa_placeholder→originale).
        Elabora le entità in ordine inverso per preservare gli indici.
        """
        sorted_ents = sorted(entities, key=lambda e: e.start, reverse=True)
        result = list(text)
        entity_map: Dict[str, str] = {}

        for ent in sorted_ents:
            token = self._get_token(ent.value, ent.type)
            entity_map[token] = ent.value
            result[ent.start:ent.end] = list(token)

        return "".join(result), entity_map

    def _get_token(self, value: str, etype: EntityType) -> str:
        key = value.strip().lower()
        if key in self._value_to_token:
            return self._value_to_token[key]
        n = self._counters.get(etype.value, 0) + 1
        self._counters[etype.value] = n
        token = f"[{etype.value}_{n:03d}]"
        self._value_to_token[key] = token
        return token


# ──────────────────────────────────────────────────────────────
# MODULO 4 – SECURE MAPPING STORE
# ──────────────────────────────────────────────────────────────

class SecureMappingStore:
    """
    Mappa locale placeholder → originale.
    Non viene mai inviata all'LLM.
    Supporta cifratura AES-256 (tramite cryptography/Fernet).

    Derivazione chiave: PBKDF2-HMAC-SHA256, 480.000 iterazioni.
    Ogni file cifrato incorpora un salt casuale di 16 byte (formato v2), così che
    la stessa passphrase produca chiavi diverse su file diversi, neutralizzando
    attacchi con tabelle precalcolate. I file legacy (salt statico) restano
    leggibili per retrocompatibilità.
    """

    _KDF_ITERATIONS = 480_000
    _MAGIC = b"PPX2"               # header dei file cifrati v2 (salt casuale)
    _SALT_LEN = 16
    _LEGACY_SALT = b"privacy_proxy_v1"   # salt statico dei file v1 (solo lettura)

    def __init__(
        self,
        store_path: Optional[str] = None,
        encryption_key: Optional[str] = None,
    ) -> None:
        self._path = Path(store_path) if store_path else None
        self._mapping: Dict[str, str] = {}
        self._passphrase: Optional[str] = None
        # ``_fernet`` resta esposto (None se cifratura non attiva) per
        # retrocompatibilità con codice/test che ne verificano la presenza.
        self._fernet = None
        if encryption_key:
            self._init_crypto(encryption_key)

    def _init_crypto(self, passphrase: str) -> None:
        self._passphrase = passphrase
        try:
            # La derivazione importa cryptography: se manca, ImportError → fallback
            # in chiaro. Il Fernet "rappresentativo" (salt statico) serve solo a
            # segnalare che la cifratura è abilitata; la chiave reale di save/load
            # è derivata per-file con salt casuale.
            self._fernet = self._derive_fernet(self._LEGACY_SALT)
        except ImportError:
            self._passphrase = None
            self._fernet = None
            logger.warning("cryptography non installato – mappa salvata in chiaro")
            return
        logger.info("Cifratura AES-256 abilitata")

    def _derive_fernet(self, salt: bytes):
        from cryptography.fernet import Fernet
        from cryptography.hazmat.primitives import hashes
        from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
        import base64

        kdf = PBKDF2HMAC(
            algorithm=hashes.SHA256(),
            length=32,
            salt=salt,
            iterations=self._KDF_ITERATIONS,
        )
        key = base64.urlsafe_b64encode(kdf.derive(self._passphrase.encode()))
        return Fernet(key)

    def store(self, entity_map: Dict[str, str]) -> None:
        self._mapping.update(entity_map)
        if self._path:
            self._save()

    def get_mapping(self) -> Dict[str, str]:
        return dict(self._mapping)

    def clear(self) -> None:
        self._mapping.clear()

    def _save(self) -> None:
        raw = json.dumps(self._mapping, ensure_ascii=False).encode()
        if self._passphrase is not None:
            import os as _os
            salt = _os.urandom(self._SALT_LEN)
            token = self._derive_fernet(salt).encrypt(raw)
            self._path.write_bytes(self._MAGIC + salt + token)
        else:
            self._path.write_text(
                json.dumps(self._mapping, ensure_ascii=False, indent=2), encoding="utf-8"
            )
        logger.info(f"Mappa salvata in {self._path}")

    def load(self) -> None:
        if not self._path or not self._path.exists():
            return
        if self._passphrase is not None:
            data = self._path.read_bytes()
            if data.startswith(self._MAGIC):
                salt = data[len(self._MAGIC):len(self._MAGIC) + self._SALT_LEN]
                token = data[len(self._MAGIC) + self._SALT_LEN:]
                fernet = self._derive_fernet(salt)
            else:
                # File legacy v1: salt statico, token Fernet "nudo".
                fernet = self._derive_fernet(self._LEGACY_SALT)
                token = data
            raw = fernet.decrypt(token)
            self._mapping = json.loads(raw.decode())
        else:
            self._mapping = json.loads(self._path.read_text(encoding="utf-8"))


# ──────────────────────────────────────────────────────────────
# MODULO 5 – LLM GATEWAY
# ──────────────────────────────────────────────────────────────

class LLMProvider(str, Enum):
    OPENAI    = "openai"
    ANTHROPIC = "anthropic"
    GEMINI    = "gemini"
    OLLAMA    = "ollama"
    DEMO      = "demo"      # Echo provider per test senza API key


class TextChunker:
    """
    Spezza un testo lungo in chunk di dimensione massima configurabile.
    Spec Modulo 6: "Gestire: documenti lunghi".

    Regole di splitting (in ordine di preferenza):
      1. Paragrafo (doppio a-capo)
      2. Frase (punto + spazio)
      3. Spazio (parola)
      4. Taglio netto se nessun separatore trovato nel limite

    Garantisce che i placeholder [TYPE_NNN] non vengano mai spezzati a metà.
    """

    _TOKEN_RE = re.compile(r"\[[A-Z_]+_[0-9]{3,}\]")

    def __init__(self, max_chars: int = 12_000) -> None:
        self.max_chars = max_chars

    def chunk(self, text: str) -> List[str]:
        if len(text) <= self.max_chars:
            return [text]

        chunks: List[str] = []
        remaining = text

        while remaining:
            if len(remaining) <= self.max_chars:
                chunks.append(remaining)
                break

            window = remaining[: self.max_chars]
            # Non tagliare dentro un placeholder [TYPE_NNN]. Un token che attraversa
            # il confine inizia prima di max_chars ma finisce dopo, quindi NON è
            # interamente contenuto in `window`: va cercato in una finestra un po'
            # più ampia (i token sono corti, < ~16 char). Se ne trova uno che
            # "scavalca" max_chars, si taglia subito prima del suo inizio.
            probe = remaining[: self.max_chars + 16]
            straddling = [
                m.start()
                for m in self._TOKEN_RE.finditer(probe)
                if m.start() < self.max_chars < m.end()
            ]
            last_token_start = min(straddling) if straddling else None
            if last_token_start is not None and last_token_start > 0:
                cut = last_token_start
            else:
                # Try paragraph boundary
                cut = window.rfind("\n\n")
                if cut <= 0:
                    cut = window.rfind(". ")
                if cut <= 0:
                    cut = window.rfind(" ")
                if cut <= 0:
                    cut = self.max_chars

            chunks.append(remaining[:cut].rstrip())
            remaining = remaining[cut:].lstrip()

        return [c for c in chunks if c]


class LLMGateway:
    """Interfaccia unificata verso OpenAI, Anthropic, Gemini, Ollama."""

    _DEFAULTS = {
        LLMProvider.OPENAI:    "gpt-4o-mini",
        LLMProvider.ANTHROPIC: "claude-haiku-4-5-20251001",
        LLMProvider.GEMINI:    "gemini-1.5-flash",
        LLMProvider.OLLAMA:    "llama3.2",
        LLMProvider.DEMO:      "echo",
    }

    def __init__(
        self,
        provider: LLMProvider,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
        base_url: str = "http://localhost:11434",
        max_chunk_chars: int = 12_000,
    ) -> None:
        self.provider = provider
        self.api_key = api_key
        self.model = model or self._DEFAULTS[provider]
        self.base_url = base_url
        self._chunker = TextChunker(max_chars=max_chunk_chars)

    def send(self, prompt: str, system_prompt: Optional[str] = None) -> str:
        """Invia un singolo prompt all'LLM (senza chunking)."""
        logger.info(f"LLM Gateway → {self.provider.value} / {self.model}")
        dispatch = {
            LLMProvider.OPENAI:    self._openai,
            LLMProvider.ANTHROPIC: self._anthropic,
            LLMProvider.GEMINI:    self._gemini,
            LLMProvider.OLLAMA:    self._ollama,
            LLMProvider.DEMO:      self._demo,
        }
        return dispatch[self.provider](prompt, system_prompt)

    @staticmethod
    def _apply_template(template: str, document: str) -> str:
        """Inserisce il documento nel template sostituendo SOLO il segnaposto
        ``{document}``.

        Si usa una replace letterale invece di ``str.format()`` perché il template
        è fornito dall'utente e può contenere altre graffe (es. richieste di output
        JSON, snippet di codice): con ``format()`` solleverebbero ``KeyError`` o
        ``ValueError``. Se ``{document}`` è assente il template resta invariato
        (comportamento storico).
        """
        return template.replace("{document}", document)

    def send_document(
        self,
        document: str,
        prompt_template: str,
        system_prompt: Optional[str] = None,
    ) -> str:
        """
        Invia un documento anonimizzato all'LLM con chunking automatico.
        Spec Modulo 6: "Gestire: documenti lunghi".

        Se il documento supera max_chunk_chars, viene spezzato in chunk,
        ogni chunk viene inviato separatamente e le risposte vengono
        concatenate in ordine.
        """
        chunks = self._chunker.chunk(document)

        if len(chunks) == 1:
            return self.send(self._apply_template(prompt_template, document), system_prompt)

        logger.info(f"Documento lungo: elaborazione in {len(chunks)} chunk…")
        responses = []
        for i, chunk in enumerate(chunks, 1):
            logger.info(f"  Chunk {i}/{len(chunks)} ({len(chunk)} char)…")
            chunk_prompt = (
                f"[Parte {i} di {len(chunks)}]\n\n"
                + self._apply_template(prompt_template, chunk)
            )
            responses.append(self.send(chunk_prompt, system_prompt))

        return "\n\n".join(responses)

    def _openai(self, prompt: str, system: Optional[str]) -> str:
        try:
            import openai
        except ImportError:
            raise ImportError("pip install openai")
        client = openai.OpenAI(api_key=self.api_key or os.environ.get("OPENAI_API_KEY"))
        msgs = []
        if system:
            msgs.append({"role": "system", "content": system})
        msgs.append({"role": "user", "content": prompt})
        resp = client.chat.completions.create(model=self.model, messages=msgs)
        # content può essere None (es. risposta filtrata o solo-tool): normalizza a "".
        return resp.choices[0].message.content or ""

    def _anthropic(self, prompt: str, system: Optional[str]) -> str:
        try:
            import anthropic
        except ImportError:
            raise ImportError("pip install anthropic")
        client = anthropic.Anthropic(api_key=self.api_key or os.environ.get("ANTHROPIC_API_KEY"))
        kwargs: Dict[str, Any] = {
            "model": self.model,
            "max_tokens": 8096,
            "messages": [{"role": "user", "content": prompt}],
        }
        if system:
            kwargs["system"] = system
        resp = client.messages.create(**kwargs)
        return resp.content[0].text

    def _gemini(self, prompt: str, system: Optional[str]) -> str:
        try:
            import google.generativeai as genai
        except ImportError:
            raise ImportError("pip install google-generativeai")
        genai.configure(api_key=self.api_key or os.environ.get("GOOGLE_API_KEY"))
        full = f"{system}\n\n{prompt}" if system else prompt
        return genai.GenerativeModel(self.model).generate_content(full).text

    @staticmethod
    def _demo(prompt: str, system: Optional[str]) -> str:
        """Provider di demo: restituisce un riassunto simulato senza API key."""
        # Estrai solo la parte {document} dal prompt per il report
        lines = prompt.strip().split("\n")
        doc_lines = [l for l in lines if l.startswith("[") or "→" in l or len(l) > 5]
        preview = "\n".join(doc_lines[:8])
        return (
            "[DEMO - nessuna API key richiesta]\n\n"
            "Il documento anonimizzato è stato elaborato. "
            "I dati sensibili sono stati sostituiti con placeholder prima dell'invio.\n\n"
            f"Anteprima testo ricevuto (prime righe):\n{preview}\n\n"
            "Per usare un LLM reale, specificare --provider openai|anthropic|gemini|ollama "
            "con la relativa API key."
        )

    def _ollama(self, prompt: str, system: Optional[str]) -> str:
        try:
            import requests
        except ImportError:
            raise ImportError("pip install requests")
        payload: Dict[str, Any] = {"model": self.model, "prompt": prompt, "stream": False}
        if system:
            payload["system"] = system
        try:
            r = requests.post(f"{self.base_url}/api/generate", json=payload, timeout=120)
            r.raise_for_status()
            return r.json()["response"]
        except Exception as exc:
            raise RuntimeError(
                f"Errore Ollama: {exc}. Verificare che Ollama sia avviato su {self.base_url}"
            ) from exc


# ──────────────────────────────────────────────────────────────
# MODULO 6 – REHYDRATION ENGINE
# ──────────────────────────────────────────────────────────────

class RehydrationEngine:
    """Ripristina i dati originali nei placeholder presenti nella risposta LLM."""

    # Case-insensitive: gli LLM a volte restituiscono token in minuscolo/misto.
    # {3,} per riconoscere anche indici a 4+ cifre (>999 entità dello stesso tipo).
    _TOKEN_RE = re.compile(r"\[[A-Za-z_]+_[0-9]{3,}\]", re.IGNORECASE)

    def rehydrate(self, text: str, entity_map: Dict[str, str]) -> str:
        """Sostituisce ogni token con il valore originale corrispondente.

        Il matching è case-insensitive per gestire LLM che alterano le maiuscole
        nei placeholder (es. [person_001] invece di [PERSON_001]).

        Implementazione a passata singola: una sola regex di alternanza di tutti i
        token, con i token ordinati dal più lungo al più corto così che, in caso di
        prefissi comuni, vinca il match più lungo. La funzione di replacement
        restituisce il valore originale LETTERALE (immune dall'interpretazione di
        sequenze come \\1 o \\g<name>, es. backslash nei percorsi Windows).
        Vantaggi rispetto a un ciclo di re.sub per-token:
          * complessità O(lunghezza_testo) invece di O(num_token × lunghezza_testo)
            → niente rallentamenti con mappe molto grandi (mitigazione DoS);
          * non ri-sostituisce token che compaiono DENTRO un valore già ripristinato.
        """
        if not entity_map:
            return text
        tokens = sorted(entity_map.keys(), key=len, reverse=True)
        pattern = re.compile("|".join(re.escape(t) for t in tokens), re.IGNORECASE)
        lookup = {k.upper(): v for k, v in entity_map.items()}
        return pattern.sub(
            lambda m: lookup.get(m.group(0).upper(), m.group(0)), text
        )

    def find_unreplaced(self, text: str) -> List[str]:
        return self._TOKEN_RE.findall(text)


# ──────────────────────────────────────────────────────────────
# MODULO 7 – VALIDATION LAYER
# ──────────────────────────────────────────────────────────────

class ValidationLayer:
    """Verifica integrità del processo di anonimizzazione e ripristino."""

    # {3,} per intercettare anche indici a 4+ cifre (>999 entità dello stesso tipo),
    # così che token non sostituiti o inventati non sfuggano alla validazione.
    _TOKEN_RE = re.compile(r"\[[A-Za-z_]+_[0-9]{3,}\]", re.IGNORECASE)

    def validate_anonymized(
        self,
        anonymized: str,
        entity_map: Dict[str, str],
    ) -> ValidationResult:
        warnings: List[str] = []
        errors: List[str] = []

        for token, original in entity_map.items():
            if original in anonymized:
                errors.append(
                    f"Dato originale trapelato nel testo anonimizzato: '{original[:30]}'"
                )

        map_upper = {k.upper() for k in entity_map}
        for tok in self._TOKEN_RE.findall(anonymized):
            if tok.upper() not in map_upper:
                errors.append(f"Token sconosciuto nel testo anonimizzato: {tok}")

        return ValidationResult(is_valid=not errors, warnings=warnings, errors=errors)

    def validate_rehydrated(
        self,
        rehydrated: str,
        entity_map: Dict[str, str],
    ) -> ValidationResult:
        warnings: List[str] = []
        errors: List[str] = []

        map_upper = {k.upper(): v for k, v in entity_map.items()}
        for tok in self._TOKEN_RE.findall(rehydrated):
            if tok.upper() in map_upper:
                errors.append(f"Token non sostituito nella risposta finale: {tok}")
            else:
                errors.append(f"Placeholder inventato dall'LLM: {tok}")

        return ValidationResult(is_valid=not errors, warnings=warnings, errors=errors)


# ──────────────────────────────────────────────────────────────
# PRIVACY PROXY – ORCHESTRATORE PRINCIPALE
# ──────────────────────────────────────────────────────────────

class PrivacyProxy:
    """
    Orchestratore della pipeline completa:
      load → detect → tokenize → store → llm → rehydrate → validate
    """

    def __init__(
        self,
        provider: LLMProvider = LLMProvider.ANTHROPIC,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
        confidence_threshold: float = 0.7,
        use_ner: bool = True,
        encryption_key: Optional[str] = None,
        mapping_store_path: Optional[str] = None,
        ollama_base_url: str = "http://localhost:11434",
        max_chunk_chars: int = 12_000,
        entity_types: Optional[List[EntityType]] = None,
    ) -> None:
        self._input       = InputLayer()
        self._detector    = SensitiveDataDetectionEngine(
            confidence_threshold, use_ner, entity_types
        )
        self._tokenizer   = TokenizationEngine()
        self._store       = SecureMappingStore(mapping_store_path, encryption_key)
        self._llm         = LLMGateway(provider, api_key, model, ollama_base_url, max_chunk_chars)
        self._rehydrator  = RehydrationEngine()
        self._validator   = ValidationLayer()

    # ── public API ───────────────────────────────────────────

    def process(
        self,
        source: str,
        prompt_template: str = "Analizza il seguente testo:\n\n{document}",
        system_prompt: Optional[str] = None,
        stop_on_error: bool = True,
    ) -> Dict[str, Any]:
        """
        Pipeline completa con LLM.

        Args:
            stop_on_error: se True (default), interrompe la pipeline prima di inviare all'LLM
                           qualora la validazione dell'anonimizzazione rilevi errori.
                           Questo garantisce che dati sensibili non trapelati vengano mai inviati.
        """

        # 1. Carica
        logger.info("[1/7] Caricamento documento…")
        doc = self._input.load(source)

        # 2. Rileva entità
        logger.info("[2/7] Rilevamento dati sensibili…")
        entities = self._detector.detect(doc.text)
        logger.info(f"      {len(entities)} entità rilevate")

        # 3. Tokenizza
        logger.info("[3/7] Anonimizzazione…")
        self._tokenizer.reset()
        anon_text, entity_map = self._tokenizer.tokenize(doc.text, entities)

        # 4. Salva mappa
        logger.info("[4/7] Salvataggio mappa sicura…")
        self._store.clear()
        self._store.store(entity_map)

        # 5. Valida anonimizzazione PRIMA di inviare all'LLM
        val_anon = self._validator.validate_anonymized(anon_text, entity_map)
        self._log_validation("Anonimizzazione", val_anon)

        if stop_on_error and not val_anon.is_valid:
            raise RuntimeError(
                "Anonimizzazione non sicura – pipeline interrotta per proteggere i dati.\n"
                f"Errori: {val_anon.errors}\n"
                "Usare --no-stop-on-error per ignorare (sconsigliato)."
            )

        # 6. Invia all'LLM (con chunking automatico per documenti lunghi)
        logger.info("[5/7] Invio testo anonimizzato all'LLM…")
        llm_resp = self._llm.send_document(anon_text, prompt_template, system_prompt)

        # 7. Ripristina
        logger.info("[6/7] Ripristino dati originali…")
        final = self._rehydrator.rehydrate(llm_resp, entity_map)

        # 8. Valida risultato finale
        logger.info("[7/7] Validazione finale…")
        val_final = self._validator.validate_rehydrated(final, entity_map)
        self._log_validation("Ripristino", val_final)

        return {
            "source": source,
            "detected_entities": [
                {
                    "value": e.value,
                    "type": e.type.value,
                    "confidence": round(e.confidence, 3),
                    "source": e.source,
                }
                for e in entities
            ],
            "anonymized_text": anon_text,
            "llm_response_anonymized": llm_resp,
            "final_response": final,
            "entity_map": entity_map,
            "validation": {
                "anonymization": {
                    "is_valid": val_anon.is_valid,
                    "warnings": val_anon.warnings,
                    "errors": val_anon.errors,
                },
                "rehydration": {
                    "is_valid": val_final.is_valid,
                    "warnings": val_final.warnings,
                    "errors": val_final.errors,
                },
            },
        }

    def anonymize_only(self, source: str) -> Dict[str, Any]:
        """Anonimizza senza inviare all'LLM."""
        doc = self._input.load(source)
        entities = self._detector.detect(doc.text)
        self._tokenizer.reset()
        anon_text, entity_map = self._tokenizer.tokenize(doc.text, entities)
        self._store.clear()
        self._store.store(entity_map)
        val = self._validator.validate_anonymized(anon_text, entity_map)
        return {
            "original_text": doc.text,
            "anonymized_text": anon_text,
            "entities": [
                {
                    "value": e.value,
                    "type": e.type.value,
                    "confidence": round(e.confidence, 3),
                    "source": e.source,
                }
                for e in entities
            ],
            "entity_map": entity_map,
            "validation": {
                "is_valid": val.is_valid,
                "warnings": val.warnings,
                "errors": val.errors,
            },
        }

    def rehydrate_only(self, text: str, entity_map: Dict[str, str]) -> str:
        """Ripristina i dati da una mappa fornita esternamente."""
        result = self._rehydrator.rehydrate(text, entity_map)
        val = self._validator.validate_rehydrated(result, entity_map)
        self._log_validation("Ripristino", val)
        return result

    # ── helpers ──────────────────────────────────────────────

    @staticmethod
    def _log_validation(label: str, v: ValidationResult) -> None:
        status = "OK" if v.is_valid else "ERRORI"
        logger.info(f"      {label}: {status}")
        for w in v.warnings:
            logger.warning(f"  ⚠  {w}")
        for e in v.errors:
            logger.error(f"  ✖  {e}")


# ──────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────

def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="privacy_proxy",
        description="Privacy Proxy per LLM – anonimizza documenti prima di inviarli agli LLM cloud",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Esempi:
  # Solo anonimizzazione (senza LLM)
  python pProxy.py --file contratto.pdf --anonymize-only --show-map

  # Analisi con Anthropic (legge ANTHROPIC_API_KEY dall'ambiente)
  python pProxy.py --file verbale.docx --provider anthropic

  # Analisi con OpenAI specificando chiave e modello
  python pProxy.py --file dati.csv --provider openai --api-key sk-... --model gpt-4o

  # Analisi con Gemini
  python pProxy.py --file report.txt --provider gemini

  # Analisi con Ollama locale
  python pProxy.py --file documento.txt --provider ollama --model llama3.2

  # Testo libero diretto
  python pProxy.py --text "Mario Rossi, CF: RSSMRA80A01H501Z, tel 333-1234567" \\
      --provider anthropic --anonymize-only

  # Con cifratura della mappa e salvataggio
  python pProxy.py --file doc.pdf --provider anthropic \\
      --encryption-key "mia-password-segreta" --save-mapping mappa.enc \\
      --output risultato.json
        """,
    )

    # ── Input ─────────────────────────────────────────────
    src = p.add_mutually_exclusive_group(required=False)
    src.add_argument("--file", "-f", metavar="PATH",
                     help="File da elaborare (PDF, CSV, TXT, DOCX, JSON)")
    src.add_argument("--text", "-t", metavar="TESTO",
                     help="Testo libero da anonimizzare")
    src.add_argument("--dir", "-d", metavar="CARTELLA",
                     help="Elabora tutti i file supportati in una cartella (batch)")

    # ── LLM ───────────────────────────────────────────────
    p.add_argument("--provider", "-p",
                   choices=["openai", "anthropic", "gemini", "ollama", "demo"],
                   default="anthropic",
                   help="Provider LLM (default: anthropic). 'demo' non richiede API key.")
    p.add_argument("--api-key", "-k", help="API key del provider")
    p.add_argument("--model", "-m", help="Modello da usare (default: dipende dal provider)")
    p.add_argument("--ollama-url", default="http://localhost:11434",
                   help="URL base Ollama (default: http://localhost:11434)")

    # ── Prompt ────────────────────────────────────────────
    p.add_argument(
        "--prompt",
        default="Analizza il seguente documento e fornisci un riassunto dettagliato:\n\n{document}",
        help="Template del prompt (usa {document} come segnaposto per il testo)",
    )
    p.add_argument("--system-prompt", help="System prompt opzionale per l'LLM")

    # ── Detection ─────────────────────────────────────────
    p.add_argument("--confidence", type=float, default=0.7, metavar="0-1",
                   help="Soglia di confidenza per il rilevamento (default: 0.7)")
    p.add_argument("--entity-types", metavar="TIPI",
                   help="Filtra i tipi da rilevare (virgola-separati, es: EMAIL,PHONE,IBAN). "
                        "Default: tutti. Valori: "
                        + ", ".join(e.value for e in EntityType))
    p.add_argument("--no-ner", action="store_true",
                   help="Usa solo regex (disabilita NER – più veloce ma meno accurato)")
    p.add_argument("--max-chunk", type=int, default=12_000, metavar="CHARS",
                   help="Dimensione massima di ogni chunk in caratteri per documenti lunghi "
                        "(default: 12000). 0 = nessun chunking.")

    # ── Security ──────────────────────────────────────────
    p.add_argument("--encryption-key", metavar="PASSPHRASE",
                   help="Cifra la mappa con AES-256 usando questa passphrase")
    p.add_argument("--save-mapping", metavar="FILE",
                   help="Salva la mappa delle entità in questo file")
    p.add_argument("--load-mapping", metavar="FILE",
                   help="Carica una mappa salvata in precedenza (usare con --rehydrate-from)")
    p.add_argument("--rehydrate-from", metavar="FILE_O_TESTO",
                   help="Testo o file con placeholder da ripristinare usando --load-mapping. "
                        "Non richiede --file né --text.")

    # ── Safety ────────────────────────────────────────────
    p.add_argument("--no-stop-on-error", action="store_true",
                   help="Non interrompere la pipeline se la validazione rileva dati trapelati "
                        "(sconsigliato: può esporre dati sensibili all'LLM)")

    # ── Output ────────────────────────────────────────────
    p.add_argument("--dry-run", action="store_true",
                   help="Rileva le entità e mostra cosa verrebbe anonimizzato, "
                        "senza modificare il testo (anteprima sicura)")
    p.add_argument("--highlight", action="store_true",
                   help="Stampa il testo originale con le entità marcate inline: "
                        "{valore|TIPO:conf} — utile per revisione visiva")
    p.add_argument("--anonymize-only", action="store_true",
                   help="Solo anonimizzazione, senza inviare all'LLM")
    p.add_argument("--redact", action="store_true",
                   help="Redazione permanente: sostituisce i dati sensibili con [REDACTED], "
                        "nessuna mappa viene conservata")
    p.add_argument("--output", "-o", metavar="FILE",
                   help="Salva il risultato completo in JSON")
    p.add_argument("--stats", action="store_true",
                   help="Stampa statistiche dettagliate sulle entità rilevate "
                        "(conteggi per tipo, distribuzione confidence, sorgente)")
    p.add_argument("--show-map", action="store_true",
                   help="Mostra la mappa placeholder → valore originale")
    p.add_argument("--show-sources", action="store_true",
                   help="Mostra il rilevatore che ha trovato ogni entità (rule/spacy/gliner/presidio)")
    p.add_argument("--verbose", "-v", action="store_true", help="Output dettagliato")

    return p


def _hr(title: str = "", width: int = 60) -> None:
    if title:
        print(f"\n{'─'*width}")
        print(f" {title}")
        print("─" * width)
    else:
        print("─" * width)


def _write_output(path: str, content: str) -> None:
    """Scrive ``content`` su ``path`` creando le cartelle intermedie se mancano.

    Centralizza tutte le scritture di output della CLI così che un percorso con
    cartelle non ancora esistenti (es. ``--output risultati/out.json``) non
    fallisca con FileNotFoundError.
    """
    out = Path(path)
    if out.parent and not out.parent.exists():
        out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(content, encoding="utf-8")


# Errori di I/O/parsing "attesi" durante il caricamento dell'input: libreria
# opzionale mancante (ImportError), formato non supportato o JSON malformato
# (JSONDecodeError ⊂ ValueError), file inesistente/illeggibile (⊂ OSError).
# Vengono convertiti in messaggi puliti dalla CLI invece di un traceback.
_INPUT_ERRORS = (ImportError, ValueError, OSError)


def _load_source_safe(source: str) -> Optional["RawTextDocument"]:
    """Carica il source gestendo gli errori attesi; ``None`` (+ messaggio) se fallisce."""
    try:
        return InputLayer().load(source)
    except _INPUT_ERRORS as exc:
        print(f"\n✖ Errore durante il caricamento dell'input: {exc}", file=sys.stderr)
        return None


def _print_entity_stats(entities: List[Dict[str, Any]]) -> None:
    """Stampa statistiche dettagliate sulle entità rilevate."""
    if not entities:
        print("  (nessuna entità rilevata)")
        return

    # Raggruppa per tipo
    by_type: Dict[str, List[Dict[str, Any]]] = {}
    for e in entities:
        by_type.setdefault(e["type"], []).append(e)

    # Raggruppa per sorgente
    by_source: Dict[str, int] = {}
    for e in entities:
        by_source[e.get("source", "rule")] = by_source.get(e.get("source", "rule"), 0) + 1

    total = len(entities)
    print(f"  Totale entità  : {total}")
    print()
    print(f"  {'Tipo':<10} {'N':>4}  {'%':>5}  {'conf.min':>8}  {'conf.max':>8}  {'conf.avg':>8}")
    print(f"  {'─'*10} {'─'*4}  {'─'*5}  {'─'*8}  {'─'*8}  {'─'*8}")
    for etype in sorted(by_type):
        group = by_type[etype]
        n = len(group)
        pct = n / total * 100
        confs = [e["confidence"] for e in group]
        print(
            f"  {etype:<10} {n:>4}  {pct:>5.1f}%"
            f"  {min(confs):>8.3f}  {max(confs):>8.3f}  {sum(confs)/len(confs):>8.3f}"
        )
    print()
    print("  Sorgente rilevamento:")
    for src, count in sorted(by_source.items(), key=lambda x: -x[1]):
        print(f"    {src:<10} {count:>4}  ({count/total*100:.1f}%)")


def _build_highlighted_text(text: str, entities: List["DetectedEntity"]) -> str:
    """Ricostruisce il testo con entità marcate: {valore|TIPO:conf}."""
    if not entities:
        return text
    parts: List[str] = []
    cursor = 0
    for ent in sorted(entities, key=lambda e: e.start):
        if ent.start < cursor:
            continue  # skip overlapping (already deduped, but be safe)
        parts.append(text[cursor:ent.start])
        parts.append(f"{{{ent.value}|{ent.type.value}:{ent.confidence:.2f}}}")
        cursor = ent.end
    parts.append(text[cursor:])
    return "".join(parts)


def _parse_entity_types(raw: Optional[str]) -> Optional[List[EntityType]]:
    """Converte stringa CSV di tipi entità in lista EntityType; None = tutti.

    Accetta sia il valore dell'enum (es. ORG, LOC, ADDR) sia il nome più esteso
    (ORGANIZATION, LOCATION, ADDRESS): per tre tipi i due differiscono e gli
    esempi del README usano entrambe le forme. I duplicati (es. 'ADDR,ADDRESS')
    e i token vuoti (virgole di troppo) vengono ignorati.
    """
    if not raw:
        return None
    valid = {e.value.upper(): e for e in EntityType}
    valid.update({e.name.upper(): e for e in EntityType})
    result: List[EntityType] = []
    seen: set = set()
    for token in raw.upper().split(","):
        token = token.strip()
        if not token:
            continue
        if token not in valid:
            canonical = ", ".join(e.value for e in EntityType)
            raise ValueError(
                f"Tipo entità sconosciuto: '{token}'. Valori validi: {canonical}"
            )
        etype = valid[token]
        if etype not in seen:
            seen.add(etype)
            result.append(etype)
    return result or None


def main(argv: Optional[List[str]] = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    _configure_logging(verbose=args.verbose)

    # Parsing filtro tipi entità
    try:
        entity_types_filter = _parse_entity_types(getattr(args, "entity_types", None))
    except ValueError as exc:
        print(f"✖ {exc}", file=sys.stderr)
        return 1

    # ── Modalità ripristino standalone (--rehydrate-from + --load-mapping) ──
    if hasattr(args, "rehydrate_from") and args.rehydrate_from:
        if not args.load_mapping:
            print("✖ --rehydrate-from richiede anche --load-mapping FILE", file=sys.stderr)
            return 1
        store = SecureMappingStore(
            store_path=args.load_mapping,
            encryption_key=args.encryption_key,
        )
        store.load()
        mapping = store.get_mapping()
        # Il sorgente può essere un file o testo diretto
        rf_path = Path(args.rehydrate_from)
        llm_text = rf_path.read_text(encoding="utf-8") if rf_path.exists() else args.rehydrate_from
        result_text = RehydrationEngine().rehydrate(llm_text, mapping)
        _hr("TESTO RIPRISTINATO")
        print(result_text)
        if args.output:
            _write_output(args.output, result_text)
        return 0

    # ── Modalità batch (--dir) ──────────────────────────────
    if hasattr(args, "dir") and args.dir:
        dir_path = Path(args.dir)
        if not dir_path.is_dir():
            print(f"✖ '{args.dir}' non è una cartella valida", file=sys.stderr)
            return 1
        supported = {".txt", ".pdf", ".csv", ".docx", ".json"}
        files = sorted(p for p in dir_path.iterdir() if p.suffix.lower() in supported)
        if not files:
            print(f"⚠ Nessun file supportato trovato in '{args.dir}'", file=sys.stderr)
            return 0
        print(f"Elaborazione batch: {len(files)} file in '{args.dir}'\n")

        # Crea la cartella di output in anticipo: altrimenti la write_text per ogni
        # file fallirebbe con FileNotFoundError se la cartella non esiste ancora.
        if args.output:
            Path(args.output).mkdir(parents=True, exist_ok=True)

        # Il proxy (e con esso il caricamento dei modelli NER) viene costruito una
        # sola volta e riusato per tutti i file: ricrearlo a ogni iterazione
        # ricaricherebbe i modelli spaCy/GLiNER, collo di bottiglia per i batch.
        # anonymize_only() azzera tokenizer e store a ogni chiamata, quindi il
        # riuso è equivalente a una nuova istanza per ciascun file.
        proxy = PrivacyProxy(
            provider=LLMProvider(args.provider),
            api_key=args.api_key,
            model=args.model,
            confidence_threshold=args.confidence,
            use_ner=not args.no_ner,
            encryption_key=args.encryption_key,
            max_chunk_chars=args.max_chunk if args.max_chunk > 0 else 10 ** 9,
            entity_types=entity_types_filter,
        )

        errors = 0
        for fp in files:
            print(f"  {'─'*50}")
            print(f"  File: {fp.name}")
            try:
                result = proxy.anonymize_only(str(fp))
                counts: Dict[str, int] = {}
                for e in result["entities"]:
                    counts[e["type"]] = counts.get(e["type"], 0) + 1
                summary = ", ".join(f"{t}:{n}" for t, n in sorted(counts.items()))
                print(f"  Entità: {len(result['entities'])}  [{summary}]")
                if args.output:
                    out = Path(args.output) / (fp.stem + "_anon.txt")
                    out.write_text(result["anonymized_text"], encoding="utf-8")
                    print(f"  Salvato: {out.name}")
            except Exception as exc:
                print(f"  ✖ Errore: {exc}", file=sys.stderr)
                errors += 1
        print(f"\nBatch completato: {len(files)-errors}/{len(files)} OK")
        return 0 if errors == 0 else 1

    source = args.file or args.text

    if not source:
        print("✖ Specificare --file, --text, --dir o --rehydrate-from",
              file=sys.stderr)
        return 1

    # Avviso precoce se manca la API key prima che la pipeline parta
    _ENV_KEY_MAP = {
        "openai":    "OPENAI_API_KEY",
        "anthropic": "ANTHROPIC_API_KEY",
        "gemini":    "GOOGLE_API_KEY",
    }
    if (
        not args.anonymize_only
        and not args.redact
        and not args.dry_run
        and not args.highlight
        and args.provider in _ENV_KEY_MAP
        and not args.api_key
        and not os.environ.get(_ENV_KEY_MAP[args.provider])
    ):
        env_var = _ENV_KEY_MAP[args.provider]
        print(
            f"⚠ Nessuna API key per '{args.provider}'.\n"
            f"  Impostare la variabile d'ambiente {env_var} oppure usare --api-key KEY\n"
            f"  Per testare senza chiave: --provider demo --anonymize-only",
            file=sys.stderr,
        )
        return 1

    # ── Modalità dry-run: solo rilevamento, testo invariato ────
    if args.dry_run:
        doc = _load_source_safe(source)
        if doc is None:
            return 1
        detector = SensitiveDataDetectionEngine(
            confidence_threshold=args.confidence,
            use_ner=not args.no_ner,
            entity_types=entity_types_filter,
        )
        entities = detector.detect(doc.text)
        _hr(f"DRY-RUN – ENTITÀ CHE VERREBBERO ANONIMIZZATE  ({len(entities)})")
        if not entities:
            print("  (nessuna entità rilevata)")
        for e in entities:
            src_tag = f"  src={e.source}" if args.show_sources else ""
            print(
                f"  [{e.type.value:8s}] conf={e.confidence:.2f}{src_tag}"
                f"  pos={e.start}-{e.end}  {e.value[:60]!r}"
            )
        if entities:
            counts_d: Dict[str, int] = {}
            for e in entities:
                counts_d[e.type.value] = counts_d.get(e.type.value, 0) + 1
            print("\n  Riepilogo: " + "  ".join(f"{t}={n}" for t, n in sorted(counts_d.items())))
        if args.output:
            import json as _json
            out_data = [
                {
                    "value": e.value,
                    "type": e.type.value,
                    "confidence": round(e.confidence, 3),
                    "source": e.source,
                    "start": e.start,
                    "end": e.end,
                }
                for e in entities
            ]
            _write_output(
                args.output, _json.dumps(out_data, ensure_ascii=False, indent=2)
            )
            print(f"\nRisultato salvato in: {args.output}")
        return 0

    # ── Modalità highlight: testo con entità marcate inline ────
    if args.highlight:
        doc = _load_source_safe(source)
        if doc is None:
            return 1
        detector = SensitiveDataDetectionEngine(
            confidence_threshold=args.confidence,
            use_ner=not args.no_ner,
            entity_types=entity_types_filter,
        )
        entities = detector.detect(doc.text)
        highlighted = _build_highlighted_text(doc.text, entities)
        _hr(f"TESTO CON ENTITÀ EVIDENZIATE  ({len(entities)} rilevate)")
        print(highlighted)
        if args.output:
            _write_output(args.output, highlighted)
            print(f"\nRisultato salvato in: {args.output}")
        return 0

    proxy = PrivacyProxy(
        provider=LLMProvider(args.provider),
        api_key=args.api_key,
        model=args.model,
        confidence_threshold=args.confidence,
        use_ner=not args.no_ner,
        encryption_key=args.encryption_key,
        mapping_store_path=args.save_mapping,
        ollama_base_url=args.ollama_url,
        max_chunk_chars=args.max_chunk if args.max_chunk > 0 else 10 ** 9,
        entity_types=entity_types_filter,
    )

    if args.redact:
        try:
            result = proxy.anonymize_only(source)
        except _INPUT_ERRORS as exc:
            print(f"\n✖ Errore durante l'elaborazione: {exc}", file=sys.stderr)
            return 1
        _TOKEN_RE_REDACT = re.compile(r"\[[A-Za-z_]+_[0-9]{3,}\]")
        redacted_text = _TOKEN_RE_REDACT.sub("[REDACTED]", result["anonymized_text"])
        _hr("TESTO REDATTO")
        print(redacted_text)
        _hr(f"ENTITÀ REDATTE  ({len(result['entities'])})")
        counts_r: Dict[str, int] = {}
        for e in result["entities"]:
            counts_r[e["type"]] = counts_r.get(e["type"], 0) + 1
        if counts_r:
            print("  " + "  ".join(f"{t}={n}" for t, n in sorted(counts_r.items())))
        if args.output:
            _write_output(args.output, redacted_text)
            print(f"\nRisultato salvato in: {args.output}")
        return 0

    if args.anonymize_only:
        try:
            result = proxy.anonymize_only(source)
        except _INPUT_ERRORS as exc:
            print(f"\n✖ Errore durante l'elaborazione: {exc}", file=sys.stderr)
            return 1

        _hr("TESTO ANONIMIZZATO")
        print(result["anonymized_text"])

        if not result["entities"] and len(result.get("original_text", "")) > 100:
            print(
                "⚠ Nessuna entità rilevata in un documento di "
                f"{len(result['original_text'])} caratteri. "
                "Verificare encoding, parsing PDF, o abbassare --confidence.",
                file=sys.stderr,
            )

        _hr(f"ENTITÀ RILEVATE  ({len(result['entities'])})")
        for e in result["entities"]:
            src_tag = f"  src={e['source']}" if args.show_sources else ""
            print(f"  [{e['type']:8s}] conf={e['confidence']:.2f}{src_tag}  {e['value'][:60]}")
        if result["entities"]:
            counts: Dict[str, int] = {}
            for e in result["entities"]:
                counts[e["type"]] = counts.get(e["type"], 0) + 1
            print("\n  Riepilogo: " + "  ".join(f"{t}={n}" for t, n in sorted(counts.items())))

        if args.stats:
            _hr("STATISTICHE DETTAGLIATE")
            _print_entity_stats(result["entities"])

        if args.show_map:
            _hr("MAPPA ENTITÀ")
            for token, original in result["entity_map"].items():
                print(f"  {token}  →  {original}")

        val = result["validation"]
        _hr("VALIDAZIONE")
        print(f"  Stato: {'✔ OK' if val['is_valid'] else '✖ ERRORI'}")
        for w in val["warnings"]:
            print(f"  ⚠ {w}")
        for e in val["errors"]:
            print(f"  ✖ {e}")

    else:
        try:
            result = proxy.process(
                source=source,
                prompt_template=args.prompt,
                system_prompt=args.system_prompt,
                stop_on_error=not args.no_stop_on_error,
            )
        except RuntimeError as exc:
            # Blocco di sicurezza della validazione (o errore runtime del provider).
            print(f"\n✖ PIPELINE INTERROTTA: {exc}", file=sys.stderr)
            return 1
        except (ImportError, ValueError, OSError) as exc:
            # Errori operativi attesi: libreria provider non installata, file di
            # input illeggibile/malformato (JSONDecodeError ⊂ ValueError,
            # FileNotFoundError ⊂ OSError), formato non supportato. Messaggio
            # pulito invece di un traceback.
            print(f"\n✖ Errore durante l'elaborazione: {exc}", file=sys.stderr)
            return 1
        except Exception as exc:
            # Errori del provider LLM non prevedibili a priori (chiave API non
            # valida, rete, rate limit, ecc.): le librerie dei provider definiscono
            # eccezioni proprie non importabili qui. Si riporta un messaggio pulito
            # (con il tipo, per diagnosi) invece di un traceback.
            print(
                f"\n✖ Errore dal provider LLM ({type(exc).__name__}): {exc}\n"
                f"  Verificare API key, modello e connettività di rete.",
                file=sys.stderr,
            )
            return 1

        _hr("RISPOSTA FINALE (dati ripristinati)")
        print(result["final_response"])

        _hr("STATISTICHE")
        print(f"  Entità rilevate   : {len(result['detected_entities'])}")
        print(f"  Anonimizzazione   : {'✔ OK' if result['validation']['anonymization']['is_valid'] else '✖ ERRORI'}")
        print(f"  Ripristino        : {'✔ OK' if result['validation']['rehydration']['is_valid'] else '✖ ERRORI'}")
        for w in result["validation"]["anonymization"]["warnings"]:
            print(f"  ⚠ {w}")
        for e in result["validation"]["rehydration"]["errors"]:
            print(f"  ✖ {e}")

        if args.stats:
            _hr("STATISTICHE DETTAGLIATE")
            _print_entity_stats(result["detected_entities"])

        if args.show_sources:
            _hr("SORGENTI RILEVAMENTO")
            for e in result["detected_entities"]:
                print(f"  [{e['type']:8s}] conf={e['confidence']:.2f}  src={e['source']}  {e['value'][:50]}")

        if args.show_map:
            _hr("MAPPA ENTITÀ")
            for token, original in result["entity_map"].items():
                print(f"  {token}  →  {original}")

    if args.output:
        _write_output(
            args.output, json.dumps(result, ensure_ascii=False, indent=2)
        )
        print(f"\nRisultato salvato in: {args.output}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
